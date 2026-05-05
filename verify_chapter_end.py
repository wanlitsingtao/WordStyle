"""验证章节后插入是否正确"""
from docx import Document
from docx.oxml.ns import qn

doc = Document('test_after_heading_result.docx')

def is_heading(para):
    p_pr = para._element.find(qn('w:pPr'))
    if p_pr is not None:
        outline_elem = p_pr.find(qn('w:outlineLvl'))
        return outline_elem is not None
    return False

answers = [i for i, p in enumerate(doc.paragraphs) if '应答' in p.text]

print(f"应答句总数: {len(answers)}")
print("\n前20个应答句位置检查:")
print("=" * 100)

for idx, i in enumerate(answers[:20]):
    prev_text = ""
    next_text = ""
    prev_type = "无"
    next_type = "无"
    
    if i > 0:
        prev_para = doc.paragraphs[i-1]
        prev_text = prev_para.text[:40]
        prev_style = prev_para.style.name if prev_para.style else "None"
        prev_type = "标题" if is_heading(prev_para) else f"正文({prev_style})"
    
    if i + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[i+1]
        next_text = next_para.text[:40]
        next_style = next_para.style.name if next_para.style else "None"
        next_type = "标题" if is_heading(next_para) else f"正文({next_style})"
    
    # 判断是否符合规则：前是正文，后是标题（或文档末尾）
    is_correct = (prev_type == "正文" and (next_type == "标题" or i == len(doc.paragraphs) - 1))
    
    status = "✓" if is_correct else "✗"
    print(f"{status} 位置{i:4d}: 前=[{prev_type}] {prev_text}")
    print(f"           后=[{next_type}] {next_text}")

# 统计正确率
correct_count = 0
for i in answers:
    prev_type = "无"
    next_type = "无"
    
    if i > 0:
        prev_para = doc.paragraphs[i-1]
        prev_type = "标题" if is_heading(prev_para) else "正文"
    
    if i + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[i+1]
        next_type = "标题" if is_heading(next_para) else "正文"
    
    if prev_type == "正文" and (next_type == "标题" or i == len(doc.paragraphs) - 1):
        correct_count += 1

print(f"\n统计:")
print(f"  总应答句数: {len(answers)}")
print(f"  符合规则数: {correct_count}")
print(f"  正确率: {correct_count/len(answers)*100:.1f}%")
