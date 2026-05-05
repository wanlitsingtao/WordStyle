"""
查找 8js3.docx 中包含"接口位置"和"ISOS"的表格
"""
from docx import Document
from docx.oxml.ns import qn

def analyze_table(table, idx):
    """详细分析表格结构"""
    print(f"\n{'='*80}")
    print(f"表格 #{idx}")
    print(f"{'='*80}")
    print(f"行数: {len(table.rows)}")
    print(f"列数: {len(table.columns)}")
    print()
    
    # 分析每个单元格的合并属性
    for i, row in enumerate(table.rows):
        print(f"行 {i+1}:")
        for j, cell in enumerate(row.cells):
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            
            text = ''.join([p.text for p in cell.paragraphs])[:30]
            
            merge_info = []
            if tcPr is not None:
                gridSpan = tcPr.find(qn('w:gridSpan'))
                if gridSpan is not None:
                    merge_info.append(f"gridSpan={gridSpan.get(qn('w:val'))}")
                
                vMerge = tcPr.find(qn('w:vMerge'))
                if vMerge is not None:
                    val = vMerge.get(qn('w:val'))
                    if val == 'continue':
                        merge_info.append("vMerge=continue(被合并)")
                    else:
                        merge_info.append("vMerge(纵向合并起始)")
            
            if merge_info:
                print(f"  [{i},{j}] {' | '.join(merge_info)}: '{text}'")
            else:
                if text.strip():
                    print(f"  [{i},{j}]: '{text}'")
        print()

# 搜索 8js3.docx
print(" 搜索 8js3.docx 中的接口表格...")
doc = Document("E:/LingMa/8js3.docx")

found_tables = []

for idx, table in enumerate(doc.tables, 1):
    # 提取所有文本
    all_text = ""
    for row in table.rows:
        for cell in row.cells:
            all_text += cell.text + " "
    
    # 检查是否包含关键词
    if "接口位置" in all_text and ("ISOS" in all_text or "FEP" in all_text):
        found_tables.append((idx, table))
        print(f"\n 找到候选表格 #{idx}")

print(f"\n\n共找到 {len(found_tables)} 个候选表格")

# 详细分析前3个候选表格
for idx, table in found_tables[:3]:
    analyze_table(table, idx)

if len(found_tables) > 3:
    print(f"\n... 还有 {len(found_tables) - 3} 个候选表格未显示")
