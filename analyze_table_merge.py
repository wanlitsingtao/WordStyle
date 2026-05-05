"""
分析表格合并单元格结构
"""
from docx import Document
from docx.oxml.ns import qn

def analyze_table_structure(table, table_name="表格"):
    """分析表格的合并结构"""
    print(f"\n{'='*60}")
    print(f"{table_name} 结构分析")
    print(f"{'='*60}")
    print(f"行数: {len(table.rows)}")
    print(f"列数: {len(table.columns)}")
    print()
    
    # 分析每个单元格的合并属性
    for i, row in enumerate(table.rows):
        print(f"第 {i+1} 行:")
        for j, cell in enumerate(row.cells):
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            
            if tcPr is not None:
                # 检查横向合并
                gridSpan = tcPr.find(qn('w:gridSpan'))
                if gridSpan is not None:
                    span_val = gridSpan.get(qn('w:val'))
                    print(f"  单元格 [{i},{j}]: gridSpan={span_val}")
                
                # 检查纵向合并
                vMerge = tcPr.find(qn('w:vMerge'))
                if vMerge is not None:
                    vMerge_val = vMerge.get(qn('w:val'))
                    if vMerge_val == 'continue':
                        print(f"  单元格 [{i},{j}]: vMerge='continue' (被合并)")
                    else:
                        print(f"  单元格 [{i},{j}]: vMerge (起始单元格)")
                
                # 如果都没有，输出文本内容（前20个字符）
                if gridSpan is None and vMerge is None:
                    text = ''.join([p.text for p in cell.paragraphs])[:20]
                    if text:
                        print(f"  单元格 [{i},{j}]: '{text}'")
            else:
                text = ''.join([p.text for p in cell.paragraphs])[:20]
                if text:
                    print(f"  单元格 [{i},{j}]: '{text}'")
        print()

def main():
    # 分析源文档
    print(" 分析源文档表格结构")
    source_doc = Document("E:/LingMa/8js3.docx")
    
    for idx, table in enumerate(source_doc.tables, 1):
        analyze_table_structure(table, f"源表格 {idx}")
        
        if idx >= 2:  # 只分析前2个表格
            print("\n⚠️  仅分析前2个表格")
            break
    
    # 分析转换后的文档（如果存在）
    import os
    converted_file = "E:/LingMa/8js3_converted.docx"
    if os.path.exists(converted_file):
        print("\n\n📄 分析转换后文档表格结构")
        converted_doc = Document(converted_file)
        
        for idx, table in enumerate(converted_doc.tables, 1):
            analyze_table_structure(table, f"转换后表格 {idx}")
            
            if idx >= 2:
                print("\n️  仅分析前2个表格")
                break

if __name__ == "__main__":
    main()
