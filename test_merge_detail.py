# -*- coding: utf-8 -*-
"""
详细测试表格横向合并功能
对比源文档和目标文档的表格结构
"""
from docx import Document
from docx.oxml.ns import qn

def analyze_table_structure(doc, doc_name):
    """分析文档中所有表格的结构"""
    print(f"\n{'='*60}")
    print(f"分析文档: {doc_name}")
    print(f"{'='*60}")
    
    for table_idx, table in enumerate(doc.tables, 1):
        rows = len(table.rows)
        cols = len(table.columns)
        
        print(f"\n表格 {table_idx}: {rows}行 × {cols}列")
        
        has_merge = False
        
        for i, row in enumerate(table.rows):
            row_info = []
            for j, cell in enumerate(row.cells):
                # 检测 gridSpan
                tc_pr = cell._element.find(qn('w:tcPr'))
                grid_span = 1
                v_merge = None
                
                if tc_pr is not None:
                    # 检查横向合并
                    grid_span_elem = tc_pr.find(qn('w:gridSpan'))
                    if grid_span_elem is not None:
                        span_val = grid_span_elem.get(qn('w:val'))
                        if span_val:
                            try:
                                grid_span = int(span_val)
                                has_merge = True
                            except ValueError:
                                pass
                    
                    # 检查纵向合并
                    v_merge_elem = tc_pr.find(qn('w:vMerge'))
                    if v_merge_elem is not None:
                        v_merge = v_merge_elem.get(qn('w:val')) or 'continue'
                        has_merge = True
                
                # 获取单元格文本（简化）
                cell_text = cell.text[:20].replace('\n', ' ') if cell.text else ''
                
                merge_info = ""
                if grid_span > 1:
                    merge_info += f"[跨{grid_span}列]"
                if v_merge:
                    merge_info += f"[vMerge:{v_merge}]"
                
                row_info.append(f"  [{i},{j}] {merge_info} '{cell_text}'")
            
            if any('[跨' in info or '[vMerge:' in info for info in row_info):
                print("  " + "-" * 50)
                for info in row_info:
                    print(info)
        
        if not has_merge:
            print("  [提示] 此表格没有合并单元格")

def compare_tables():
    """对比源文档和目标文档的表格结构"""
    source_file = "E:/LingMa/8js3.docx"
    target_file = "E:/LingMa/test_horizontal_merge_result.docx"
    
    print("=" * 60)
    print("对比源文档和目标文档的表格结构")
    print("=" * 60)
    
    source_doc = Document(source_file)
    target_doc = Document(target_file)
    
    # 分析源文档
    analyze_table_structure(source_doc, "源文档 (8js3.docx)")
    
    # 分析目标文档
    analyze_table_structure(target_doc, "目标文档 (test_horizontal_merge_result.docx)")
    
    print(f"\n{'='*60}")
    print("对比总结")
    print(f"{'='*60}")
    print(f"源文档表格数: {len(source_doc.tables)}")
    print(f"目标文档表格数: {len(target_doc.tables)}")
    
    if len(source_doc.tables) == len(target_doc.tables):
        print("✅ 表格数量一致")
    else:
        print("❌ 表格数量不一致！")

if __name__ == "__main__":
    compare_tables()
