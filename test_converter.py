# -*- coding: utf-8 -*-
"""
测试脚本 - 验证文档转换器核心功能
"""
import os
import sys
import io

# 设置标准输出编码为UTF-8
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from doc_converter import DocumentConverter


def create_test_documents():
    """创建测试用的源文档和模板文档"""
    
    # 创建源文档
    print("创建测试源文档...")
    source_doc = Document()
    
    # 添加标题
    heading1 = source_doc.add_heading('第一章 项目概述', level=1)
    source_doc.add_paragraph('本项目必须在规定时间内完成。投标人应当具备相关资质。')
    
    heading2 = source_doc.add_heading('1.1 项目背景', level=2)
    source_doc.add_paragraph('应充分理解项目需求。不得忽视任何细节。')
    
    # 添加正文
    source_doc.add_paragraph('本项目的实施不应影响现有系统运行。切勿擅自修改配置。')
    
    # 添加列表
    list_para = source_doc.add_paragraph('项目要求：', style='List Paragraph')
    list_para.add_run('\n• 必须按时完成')
    list_para.add_run('\n• 应保证质量')
    
    # 添加表格
    table = source_doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '序号'
    table.cell(0, 1).text = '内容'
    table.cell(1, 0).text = '1'
    table.cell(1, 1).text = '投标人须提供完整方案'
    table.cell(2, 0).text = '2'
    table.cell(2, 1).text = '应符合技术规范'
    
    source_doc.save('test_source.docx')
    print("[OK] 源文档创建完成: test_source.docx")

    # 创建模板文档
    print("\n创建测试模板文档...")
    template_doc = Document()

    # 定义样式（实际使用时模板应有预定义的样式）
    template_doc.add_heading('模板标题示例', level=1)
    template_doc.add_heading('模板二级标题', level=2)
    para = template_doc.add_paragraph('模板正文样式示例')
    para.style = 'Normal'

    # 添加应答句样式
    from docx.enum.style import WD_STYLE_TYPE
    try:
        answer_style = template_doc.styles.add_style('应答句', WD_STYLE_TYPE.PARAGRAPH)
        answer_style.base_style = template_doc.styles['Normal']
    except:
        pass

    template_doc.save('test_template.docx')
    print("[OK] 模板文档创建完成: test_template.docx")


def test_conversion():
    """测试转换功能"""
    print("\n" + "="*60)
    print("开始测试文档转换功能")
    print("="*60)
    
    converter = DocumentConverter()
    
    # 测试1：样式转换
    print("\n[测试1] 样式转换...")
    success, msg = converter.convert_styles(
        'test_source.docx',
        'test_template.docx',
        'test_output_styles.docx'
    )
    status = "[OK] 成功" if success else "[FAIL] 失败"
    print(f"结果: {status}")
    print(f"消息: {msg}")

    # 测试2：语气转换
    print("\n[测试2] 祈使语气转换...")
    success, msg = converter.convert_mood(
        'test_output_styles.docx',
        'test_output_mood.docx'
    )
    status = "[OK] 成功" if success else "[FAIL] 失败"
    print(f"结果: {status}")
    print(f"消息: {msg}")

    # 测试3：插入应答句
    print("\n[测试3] 插入应答句...")
    success, msg = converter.insert_response_after_headings(
        'test_output_mood.docx',
        'test_output_final.docx'
    )
    status = "[OK] 成功" if success else "[FAIL] 失败"
    print(f"结果: {status}")
    print(f"消息: {msg}")

    # 测试4：完整转换流程
    print("\n[测试4] 完整转换流程...")

    def progress_callback(step, message):
        print(f"  进度 [{step}/7]: {message}")

    success, msg = converter.full_convert(
        source_file='test_source.docx',
        template_file='test_template.docx',
        output_file='test_complete_output.docx',
        do_mood=True,
        answer_text='应答：本投标人理解并满足要求。',
        answer_style='应答句',
        progress_callback=progress_callback
    )
    status = "[OK] 成功" if success else "[FAIL] 失败"
    print(f"\n结果: {status}")
    print(f"消息: {msg}")

    # 验证输出文件
    if os.path.exists('test_complete_output.docx'):
        print("\n[OK] 输出文件已生成: test_complete_output.docx")

        # 读取并显示部分内容
        doc = Document('test_complete_output.docx')
        print(f"\n输出文档统计:")
        print(f"  段落数: {len(doc.paragraphs)}")
        print(f"  表格数: {len(doc.tables)}")

        # 显示前几个段落
        print(f"\n前5个段落预览:")
        for i, para in enumerate(doc.paragraphs[:5], 1):
            text = para.text[:50] if para.text else "[空]"
            style = para.style.name if para.style else "无样式"
            print(f"  {i}. [{style}] {text}")
    else:
        print("\n[FAIL] 输出文件未生成")


def cleanup():
    """清理测试文件"""
    test_files = [
        'test_source.docx',
        'test_template.docx',
        'test_output_styles.docx',
        'test_output_mood.docx',
        'test_output_final.docx',
        'test_complete_output.docx'
    ]
    
    print("\n" + "="*60)
    print("清理测试文件")
    print("="*60)
    
    for f in test_files:
        if os.path.exists(f):
            try:
                os.remove(f)
                print(f"[OK] 已删除: {f}")
            except Exception as e:
                print(f"[FAIL] 删除失败 {f}: {e}")


if __name__ == "__main__":
    try:
        # 创建测试文档
        create_test_documents()

        # 执行转换测试
        test_conversion()

        print("\n" + "="*60)
        print("所有测试完成！")
        print("="*60)

    except Exception as e:
        print(f"\n[ERROR] 测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

    finally:
        # 自动清理测试文件
        print("\n正在清理测试文件...")
        cleanup()
