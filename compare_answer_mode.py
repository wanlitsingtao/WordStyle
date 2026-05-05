"""
对比测试：章节后插入 vs 示例文档
"""
from docx import Document
from docx.oxml.ns import qn

def compare_documents():
    """对比生成的文档和示例文档"""
    
    # 加载示例文档
    example_doc = Document("章节后插入应答句.docx")
    
    # 加载生成的文档
    generated_doc = Document("test_after_heading_result.docx")
    
    print("=" * 80)
    print("对比分析")
    print("=" * 80)
    
    # 分析示例文档的应答句位置
    print("\n【示例文档】应答句位置规律：")
    example_answers = []
    for i, para in enumerate(example_doc.paragraphs):
        if "应答" in para.text or "【答】" in para.text:
            example_answers.append(i)
            
            # 检查前一段和后一段
            prev_type = "无"
            next_type = "无"
            
            if i > 0:
                prev_para = example_doc.paragraphs[i-1]
                p_pr = prev_para._element.find(qn('w:pPr'))
                if p_pr is not None:
                    outline_elem = p_pr.find(qn('w:outlineLvl'))
                    if outline_elem is not None:
                        prev_type = f"标题{outline_elem.get(qn('w:val'), 0)}"
                    else:
                        prev_type = "正文"
                else:
                    prev_type = "正文"
            
            if i + 1 < len(example_doc.paragraphs):
                next_para = example_doc.paragraphs[i+1]
                p_pr = next_para._element.find(qn('w:pPr'))
                if p_pr is not None:
                    outline_elem = p_pr.find(qn('w:outlineLvl'))
                    if outline_elem is not None:
                        next_type = f"标题{outline_elem.get(qn('w:val'), 0)}"
                    else:
                        next_type = "正文"
                else:
                    next_type = "正文"
            
            print(f"  位置 {i:3d}: 前={prev_type:6s}, 后={next_type:6s}")
    
    # 分析生成文档的应答句位置
    print("\n【生成文档】应答句位置规律：")
    generated_answers = []
    for i, para in enumerate(generated_doc.paragraphs):
        if "应答" in para.text or "【答】" in para.text:
            generated_answers.append(i)
            
            # 检查前一段和后一段
            prev_type = "无"
            next_type = "无"
            
            if i > 0:
                prev_para = generated_doc.paragraphs[i-1]
                p_pr = prev_para._element.find(qn('w:pPr'))
                if p_pr is not None:
                    outline_elem = p_pr.find(qn('w:outlineLvl'))
                    if outline_elem is not None:
                        prev_type = f"标题{outline_elem.get(qn('w:val'), 0)}"
                    else:
                        prev_type = "正文"
                else:
                    prev_type = "正文"
            
            if i + 1 < len(generated_doc.paragraphs):
                next_para = generated_doc.paragraphs[i+1]
                p_pr = next_para._element.find(qn('w:pPr'))
                if p_pr is not None:
                    outline_elem = p_pr.find(qn('w:outlineLvl'))
                    if outline_elem is not None:
                        next_type = f"标题{outline_elem.get(qn('w:val'), 0)}"
                    else:
                        next_type = "正文"
                else:
                    next_type = "正文"
            
            print(f"  位置 {i:3d}: 前={prev_type:6s}, 后={next_type:6s}")
    
    print(f"\n统计：")
    print(f"  示例文档应答句数量：{len(example_answers)}")
    print(f"  生成文档应答句数量：{len(generated_answers)}")
    
    # 检查规律是否一致
    print(f"\n规律检查：")
    
    # 检查示例文档的规律
    example_pattern_correct = True
    for pos in example_answers[:5]:
        if pos + 1 < len(example_doc.paragraphs):
            next_para = example_doc.paragraphs[pos + 1]
            p_pr = next_para._element.find(qn('w:pPr'))
            if p_pr is not None:
                outline_elem = p_pr.find(qn('w:outlineLvl'))
                if outline_elem is None:
                    # 下一个不是标题，说明规律不对
                    example_pattern_correct = False
                    break
    
    # 检查生成文档的规律
    generated_pattern_correct = True
    for pos in generated_answers[:5]:
        if pos + 1 < len(generated_doc.paragraphs):
            next_para = generated_doc.paragraphs[pos + 1]
            p_pr = next_para._element.find(qn('w:pPr'))
            if p_pr is not None:
                outline_elem = p_pr.find(qn('w:outlineLvl'))
                if outline_elem is None:
                    # 下一个不是标题，说明规律不对
                    generated_pattern_correct = False
                    break
    
    print(f"  示例文档规律（应答句后是标题）：{'✓ 符合' if example_pattern_correct else '✗ 不符合'}")
    print(f"  生成文档规律（应答句后是标题）：{'✓ 符合' if generated_pattern_correct else '✗ 不符合'}")

if __name__ == "__main__":
    compare_documents()
