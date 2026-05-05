"""
精确搜索包含"至 ISOS FEP"的表格
"""
from docx import Document
from docx.oxml.ns import qn

def analyze_table_detail(table, idx):
    """详细分析表格的每一行和合并情况"""
    print(f"\n{'='*80}")
    print(f"表格 #{idx} - 详细分析")
    print(f"{'='*80}")
    print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")
    print()
    
    # 显示每一行的内容
    for i, row in enumerate(table.rows[:8]):  # 只显示前8行（通常是表头）
        print(f"行 {i+1}:")
        cells_info = []
        for j, cell in enumerate(row.cells):
            text = ''.join([p.text for p in cell.paragraphs]).strip()
            if text:
                cells_info.append(f"[{j}]:'{text[:15]}'")
        
        if cells_info:
            print(f"  内容: {' | '.join(cells_info)}")
        
        # 检查合并属性
        for j, cell in enumerate(row.cells):
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                gridSpan = tcPr.find(qn('w:gridSpan'))
                vMerge = tcPr.find(qn('w:vMerge'))
                
                if gridSpan is not None:
                    print(f"  [{i},{j}]:  横向合并 gridSpan={gridSpan.get(qn('w:val'))}")
                if vMerge is not None:
                    val = vMerge.get(qn('w:val'))
                    if val == 'continue':
                        print(f"  [{i},{j}]: ⬇️ 纵向合并(被合并)")
                    else:
                        print(f"  [{i},{j}]: ️ 纵向合并(起始)")
        print()

# 搜索 8js3.docx
print("🔍 精确搜索包含'至 ISOS FEP'的表格...\n")
doc = Document("E:/LingMa/8js3.docx")

for idx, table in enumerate(doc.tables, 1):
    all_text = ""
    for row in table.rows:
        for cell in row.cells:
            all_text += cell.text + " "
    
    # 精确匹配
    if "至 ISOS FEP" in all_text and "IBP" in all_text:
        print(f"✅ 找到目标表格 #{idx}!")
        analyze_table_detail(table, idx)
        print("\n" + "="*80)
        print("这是你要的表格吗？")
        print("="*80)
        break
