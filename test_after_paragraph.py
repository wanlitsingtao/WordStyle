"""
测试逐段后应答模式
"""
from doc_converter import DocumentConverter
import os

def test_after_paragraph_mode():
    """测试逐段后应答模式"""
    
    converter = DocumentConverter()
    
    # 使用现有的测试文件
    input_file = "8js3.docx"
    output_file = "test_after_paragraph_result.docx"
    
    if not os.path.exists(input_file):
        print(f"错误：找不到测试文件 {input_file}")
        return
    
    print("=" * 60)
    print("测试：逐段后应答")
    print("=" * 60)
    
    # 测试逐段后应答模式
    success, actual_file, msg = converter.insert_response_after_headings(
        input_file=input_file,
        output_file=output_file,
        answer_text="【应答】",
        answer_style="应答句",
        mode='after_paragraph'  # 关键参数：逐段后应答
    )
    
    if success:
        print(f"\n✅ 成功！")
        print(f"输出文件: {actual_file}")
        print(f"消息: {msg}")
        
        # 分析结果
        from docx import Document
        doc = Document(actual_file)
        
        # 统计应答句数量
        answer_count = sum(1 for para in doc.paragraphs if '应答' in para.text)
        print(f"\n应答句总数: {answer_count}")
        
        # 检查前几个应答句的位置
        print("\n前10个应答句位置检查:")
        print("-" * 80)
        count = 0
        for i, para in enumerate(doc.paragraphs):
            if '应答' in para.text and count < 10:
                prev_text = ""
                next_text = ""
                
                if i > 0:
                    prev_text = doc.paragraphs[i-1].text[:40]
                
                if i + 1 < len(doc.paragraphs):
                    next_text = doc.paragraphs[i+1].text[:40]
                
                print(f"位置{i}: 前=[{prev_text}], 后=[{next_text}]")
                count += 1
        
        print("\n" + "=" * 60)
        print("测试完成！请打开文档检查结果。")
        print("=" * 60)
    else:
        print(f"\n❌ 失败: {msg}")

if __name__ == "__main__":
    test_after_paragraph_mode()
