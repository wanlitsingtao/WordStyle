# -*- coding: utf-8 -*-
"""定位 8js3.docx 中特定章节的表格"""
from docx import Document
from docx.oxml.ns import qn

def analyze_table_detail(table, idx, doc_name):
    """详细分析表格结构"""
    print("\n" + "="*80)
    print(f"{doc_name} - 表格 #{idx}")
    print("="*80)
    print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")
    print()
    
    # 显示前8行（表头和数据行）
    for i, row in enumerate(table.rows[:8]):
        print(f"行 {i+1}:")
        for j, cell in enumerate(row.cells):
            text = ''.join([p.text for p in cell.paragraphs]).strip()
            if text:
                print(f"  [{i},{j}]: '{text[:25]}'")
        
        # 检查合并属性
        for j, cell in enumerate(row.cells):
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                gridSpan = tcPr.find(qn('w:gridSpan'))
                vMerge = tcPr.find(qn('w:vMerge'))
                
                if gridSpan is not None:
                    print(f"  -> [{i},{j}]: 横向合并 gridSpan={gridSpan.get(qn('w:val'))}")
                if vMerge is not None:
                    val = vMerge.get(qn('w:val'))
                    if val == 'continue':
                        print(f"  -> [{i},{j}]: 纵向合并(被合并)")
                    else:
                        print(f"  -> [{i},{j}]: 纵向合并(起始)")
        print()

# 分析源文档
print("分析源文档: 8js3.docx")
print("定位章节 3.8.7、3.9.7 等的设计要求表格")
print()

source_doc = Document("E:/LingMa/8js3.docx")

# 由于文档有308个表格，让我搜索包含特定关键词的表格
# 从你提供的图片，表格包含: "接口位置"、"监控要求描述"、"至 ISOS FEP"、"IBP"
# 以及子表头: "AI"、"AO"、"DI"、"DO"

keywords_to_find = ["接口位置", "监控要求描述"]

found_tables = []

for idx, table in enumerate(source_doc.tables, 1):
    all_text = ""
    for row in table.rows:
        for cell in row.cells:
            all_text += cell.text + " "
    
    # 检查是否包含关键词
    if all(kw in all_text for kw in keywords_to_find):
        # 进一步检查是否包含"AI"、"AO"等子表头
        if "AI" in all_text and "AO" in all_text:
            found_tables.append((idx, table))

print(f"找到 {len(found_tables)} 个候选表格\n")

# 分析前3个表格
for i, (idx, table) in enumerate(found_tables[:3]):
    analyze_table_detail(table, idx, "源文档")

if len(found_tables) > 3:
    print(f"\n... 还有 {len(found_tables) - 3} 个候选表格")

# 如果找到了，也分析转换后的文档
import os
converted_file = "E:/LingMa/8js3_converted.docx"
if os.path.exists(converted_file) and found_tables:
    print("\n\n" + "="*80)
    print("分析转换后文档: 8js3_converted.docx")
    print("="*80)
    
    converted_doc = Document(converted_file)
    
    # 分析对应位置的表格
    for i, (idx, source_table) in enumerate(found_tables[:3]):
        if idx <= len(converted_doc.tables):
            converted_table = converted_doc.tables[idx - 1]  # 表格索引从1开始
            analyze_table_detail(converted_table, idx, "转换后文档")
        else:
            print(f"\n表格 #{idx} 在转换后文档中不存在！")
