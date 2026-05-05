# -*- coding: utf-8 -*-
"""
测试 Visio 图和 OLE 对象处理功能
"""
import os
from docx import Document
from docx.oxml.ns import qn

def test_visio_detection():
    """测试检测 Word 文档中的 Visio 图和 OLE 对象"""
    
    # 查找所有 docx 文件
    docx_files = [f for f in os.listdir('e:\\LingMa') if f.endswith('.docx')]
    
    print("正在检查以下文件中的 Visio 图和 OLE 对象:")
    print("=" * 60)
    
    for docx_file in docx_files[:3]:  # 只检查前3个文件
        filepath = os.path.join('e:\\LingMa', docx_file)
        print(f"\n文件: {docx_file}")
        
        try:
            doc = Document(filepath)
            
            # 检查段落中的特殊对象
            object_count = 0
            shape_count = 0
            
            for para in doc.paragraphs:
                # 使用正确的命名空间
                objects = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
                shapes = para._element.findall('.//{urn:schemas-microsoft-com:vml}shape')
                
                if objects:
                    object_count += len(objects)
                    print(f"  - 发现 {len(objects)} 个 OLE 对象")
                
                if shapes:
                    shape_count += len(shapes)
                    print(f"  - 发现 {len(shapes)} 个 VML 形状")
            
            # 检查表格单元格中的特殊对象
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            objects = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
                            shapes = para._element.findall('.//{urn:schemas-microsoft-com:vml}shape')
                            
                            if objects:
                                object_count += len(objects)
                                print(f"  - 表格中发现 {len(objects)} 个 OLE 对象")
                            
                            if shapes:
                                shape_count += len(shapes)
                                print(f"  - 表格中发现 {len(shapes)} 个 VML 形状")
            
            if object_count == 0 and shape_count == 0:
                print("  - 未发现 Visio 图或 OLE 对象")
            else:
                print(f"  总计: {object_count} 个 OLE 对象, {shape_count} 个 VML 形状")
        
        except Exception as e:
            print(f"  错误: {e}")
    
    print("\n" + "=" * 60)
    print("检测完成!")

if __name__ == "__main__":
    test_visio_detection()
