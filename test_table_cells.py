"""
测试并修复表格合并单元格复制

问题分析：
当前代码在第二步（复制内容）时，会跳过 vMerge='continue' 的单元格，
但这会导致那些单元格的文本内容丢失。

正确的做法应该是：
1. 复制所有合并属性
2. 只在一个起始单元格中填充内容
3. 跳过 continue 单元格（因为它们不应该有内容）

但是，如果代码在遍历时，cell.paragraphs 会跳过被合并的单元格，
这可能导致内容映射错误。

让我检查实际的单元格遍历情况
"""
from docx import Document
from docx.oxml.ns import qn

doc = Document("E:/LingMa/8js1.docx")
table = doc.tables[0]

print("分析表格单元格遍历:")
print(f"总行数: {len(table.rows)}")
print(f"总列数: {len(table.columns)}")
print()

# 检查每一行的 cells 数量
for i, row in enumerate(table.rows[:5]):
    print(f"行 {i+1}: 有 {len(row.cells)} 个 cells")
    for j, cell in enumerate(row.cells):
        text = ''.join([p.text for p in cell.paragraphs]).strip()[:15]
        
        # 检查合并属性
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        merge_info = ""
        if tcPr is not None:
            vMerge = tcPr.find(qn('w:vMerge'))
            if vMerge is not None:
                val = vMerge.get(qn('w:val'))
                merge_info = f" [vMerge={val}]"
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is not None:
                merge_info += f" [gridSpan={gridSpan.get(qn('w:val'))}]"
        
        if text or merge_info:
            print(f"  cell[{j}]: '{text}'{merge_info}")
    print()
