"""
分析你提供的图片中的表格结构

从你的图片中，我看到的表格结构是：

【转换前】
第1行：接口位置 | 监控要求描述 | 至ISOS FEP(5列) | IBP(2列) | 备注
第2行：AI | AO | DI | DO | (智慧运行设备室) | (消防水泵) | (空白)

【关键特征】
- 有多层表头（至少2行）
- 有横向合并（gridSpan）
- 有纵向合并（vMerge）
- 第一列"接口位置"在数据行中是纵向合并的

让我尝试搜索所有包含"DI"、"DO"、"AI"、"AO"这些关键词的表格
"""
from docx import Document
from docx.oxml.ns import qn

def find_table_by_keywords(doc_file, keywords_list):
    """通过多组关键词查找表格"""
    doc = Document(doc_file)
    
    for idx, table in enumerate(doc.tables, 1):
        all_text = ""
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + " "
        
        # 检查是否包含任意一组关键词
        for keywords in keywords_list:
            if all(kw in all_text for kw in keywords):
                print(f"\n{'='*80}")
                print(f"📄 {doc_file} - 表格 #{idx}")
                print(f"{'='*80}")
                print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")
                print(f"匹配关键词: {keywords}")
                print()
                
                # 显示前5行
                for i, row in enumerate(table.rows[:5]):
                    print(f"行 {i+1}:")
                    for j, cell in enumerate(row.cells):
                        text = ''.join([p.text for p in cell.paragraphs]).strip()
                        if text:
                            print(f"  [{i},{j}]: '{text[:20]}'")
                    
                    # 检查合并属性
                    for j, cell in enumerate(row.cells):
                        tc = cell._element
                        tcPr = tc.find(qn('w:tcPr'))
                        if tcPr is not None:
                            gridSpan = tcPr.find(qn('w:gridSpan'))
                            vMerge = tcPr.find(qn('w:vMerge'))
                            
                            if gridSpan is not None:
                                print(f"  → [{i},{j}]: 横向合并 gridSpan={gridSpan.get(qn('w:val'))}")
                            if vMerge is not None:
                                val = vMerge.get(qn('w:val'))
                                print(f"  → [{i},{j}]: 纵向合并 vMerge={val}")
                    print()
                
                return True
    
    return False

# 搜索关键词组合
keyword_groups = [
    ["DI", "DO", "AI", "AO"],
    ["消防水泵", "喷淋泵"],
    ["智慧运行设备室"],
]

for doc_file in ["8js1.docx", "8js2.docx", "8js3.docx"]:
    print(f"\n🔍 搜索 {doc_file}...")
    if find_table_by_keywords(doc_file, keyword_groups):
        print(f"\n✅ 在 {doc_file} 中找到匹配表格！")
        break
