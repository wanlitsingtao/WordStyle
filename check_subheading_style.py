"""检查小标题的样式"""
from docx import Document

doc = Document('test_after_heading_result.docx')

# 查看位置2附近的段落
print("位置2附近的段落:")
for i in range(max(0, 2-2), min(len(doc.paragraphs), 2+3)):
    para = doc.paragraphs[i]
    style_name = para.style.name if para.style else "None"
    text = para.text[:60]
    marker = "[应答]" if "应答" in text else ""
    print(f"  {i}: [{style_name:20s}] {marker:8s} {text}")

print("\n\n位置18附近的段落:")
for i in range(max(0, 18-2), min(len(doc.paragraphs), 18+3)):
    para = doc.paragraphs[i]
    style_name = para.style.name if para.style else "None"
    text = para.text[:60]
    marker = "[应答]" if "应答" in text else ""
    print(f"  {i}: [{style_name:20s}] {marker:8s} {text}")
