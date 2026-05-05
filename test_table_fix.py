# -*- coding: utf-8 -*-
"""测试表格合并单元格复制修复"""
from doc_converter import StyleConverter

print("="*80)
print("测试表格合并单元格复制")
print("="*80)

# 创建转换器
converter = StyleConverter()

# 测试文件
source_file = "E:/LingMa/8js1.docx"
template_file = "E:/LingMa/mb.docx"
output_file = "E:/LingMa/test_table_merge_fix.docx"

print(f"\n源文件: {source_file}")
print(f"模板文件: {template_file}")
print(f"输出文件: {output_file}")
print()

# 执行转换
try:
    success, actual_file, msg = converter.full_convert(
        source_file=source_file,
        template_file=template_file,
        output_file=output_file,
        do_mood=False,  # 不测试语气转换，只测试表格
        do_answer_insertion=False  # 不插入应答句
    )
    
    if success:
        print(f"\n✅ 转换成功！")
        print(f"实际输出文件: {actual_file}")
        print(f"消息: {msg}")
        
        # 分析转换后的表格
        from docx import Document
        from docx.oxml.ns import qn
        
        print("\n" + "="*80)
        print("分析转换后的表格结构")
        print("="*80)
        
        converted_doc = Document(actual_file)
        
        # 分析第一个表格
        if converted_doc.tables:
            table = converted_doc.tables[0]
            print(f"\n转换后的表格 #1:")
            print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")
            print()
            
            # 检查前5行的合并属性
            for i, row in enumerate(table.rows[:5]):
                print(f"行 {i+1}:")
                for j, cell in enumerate(row.cells):
                    text = ''.join([p.text for p in cell.paragraphs]).strip()[:15]
                    
                    tc = cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    merge_info = ""
                    if tcPr is not None:
                        vMerge = tcPr.find(qn('w:vMerge'))
                        if vMerge is not None:
                            val = vMerge.get(qn('w:val'))
                            merge_info = f" [vMerge={val}]"
                        gridSpan = tcPr.find(qn('w:gridSpan'))
                        if gridSpan is not None:
                            merge_info += f" [gridSpan={gridSpan.get(qn('w:val'))}]"
                    
                    if text or merge_info:
                        print(f"  cell[{j}]: '{text}'{merge_info}")
                print()
        
        print("\n" + "="*80)
        print("请打开转换后的文档，检查表格合并是否正确！")
        print("="*80)
    else:
        print(f"\n❌ 转换失败: {msg}")

except Exception as e:
    print(f"\n❌ 发生错误: {e}")
    import traceback
    traceback.print_exc()
