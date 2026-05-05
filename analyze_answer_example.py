"""
分析"章节后插入应答句.docx"的结构
"""
from docx import Document
from docx.oxml.ns import qn

def analyze_document():
    doc = Document("章节后插入应答句.docx")
    
    print("=" * 80)
    print("文档结构分析")
    print("=" * 80)
    
    for i, para in enumerate(doc.paragraphs[:50]):  # 只看前50段
        # 获取样式名称
        style_name = para.style.name if para.style else "None"
        
        # 检查是否为标题
        is_heading = False
        outline_level = 0
        
        p_pr = para._element.find(qn('w:pPr'))
        if p_pr is not None:
            outline_elem = p_pr.find(qn('w:outlineLvl'))
            if outline_elem is not None:
                is_heading = True
                outline_level = int(outline_elem.get(qn('w:val'), 0))
        
        # 检查是否包含"应答"
        has_answer = "应答" in para.text or "【答】" in para.text
        
        # 显示信息
        marker = ""
        if is_heading:
            marker = f"[标题{outline_level}]"
        elif has_answer:
            marker = "[应答句]"
        
        text_preview = para.text[:60] if para.text else "(空段落)"
        print(f"{i:3d}. {marker:10s} | {style_name:20s} | {text_preview}")
    
    print("\n" + "=" * 80)
    print("查找应答句的位置规律...")
    print("=" * 80)
    
    # 找出所有应答句及其前后文
    answer_positions = []
    for i, para in enumerate(doc.paragraphs):
        if "应答" in para.text or "【答】" in para.text:
            answer_positions.append(i)
    
    print(f"\n共找到 {len(answer_positions)} 个应答句")
    print("\n应答句位置分析（前后各2段）：")
    
    for pos in answer_positions[:10]:  # 只看前10个
        print(f"\n--- 应答句位置: 第 {pos} 段 ---")
        
        # 前2段
        for j in range(max(0, pos-2), pos):
            prev_para = doc.paragraphs[j]
            prev_style = prev_para.style.name if prev_para.style else "None"
            prev_text = prev_para.text[:50] if prev_para.text else "(空)"
            print(f"  前{pos-j}段 [{j}]: {prev_style:20s} | {prev_text}")
        
        # 当前段（应答句）
        curr_para = doc.paragraphs[pos]
        curr_style = curr_para.style.name if curr_para.style else "None"
        print(f"  当前  [{pos}]: {curr_style:20s} | {curr_para.text[:50]}")
        
        # 后2段
        for j in range(pos+1, min(len(doc.paragraphs), pos+3)):
            next_para = doc.paragraphs[j]
            next_style = next_para.style.name if next_para.style else "None"
            next_text = next_para.text[:50] if next_para.text else "(空)"
            print(f"  后{j-pos}段 [{j}]: {next_style:20s} | {next_text}")

if __name__ == "__main__":
    analyze_document()
