# -*- coding: utf-8 -*-
"""
验证横向合并单元格是否正确保留
"""
from docx import Document
from docx.oxml.ns import qn

def check_grid_span_in_table(table, table_idx):
    """检查表格中的 gridSpan 属性"""
    grid_spans_found = []
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            tc_pr = cell._element.find(qn('w:tcPr'))
            if tc_pr is not None:
                grid_span_elem = tc_pr.find(qn('w:gridSpan'))
                if grid_span_elem is not None:
                    span_val = grid_span_elem.get(qn('w:val'))
                    if span_val:
                        try:
                            span = int(span_val)
                            if span > 1:
                                grid_spans_found.append({
                                    'row': i,
                                    'col': j,
                                    'span': span,
                                    'text': cell.text[:30]
                                })
                        except ValueError:
                            pass
    
    return grid_spans_found

def compare_horizontal_merge():
    """对比源文档和目标文档的横向合并"""
    source_file = "E:/LingMa/8js3.docx"
    target_file = "E:/LingMa/test_horizontal_merge_result.docx"
    
    print("=" * 70)
    print("对比横向合并单元格")
    print("=" * 70)
    
    source_doc = Document(source_file)
    target_doc = Document(target_file)
    
    print(f"\n源文档表格数: {len(source_doc.tables)}")
    print(f"目标文档表格数: {len(target_doc.tables)}")
    
    # 检查每个表格的横向合并
    source_has_merge = 0
    target_has_merge = 0
    match_count = 0
    mismatch_count = 0
    
    for idx in range(min(len(source_doc.tables), len(target_doc.tables))):
        source_table = source_doc.tables[idx]
        target_table = target_doc.tables[idx]
        
        source_merges = check_grid_span_in_table(source_table, idx)
        target_merges = check_grid_span_in_table(target_table, idx)
        
        if source_merges or target_merges:
            print(f"\n{'='*70}")
            print(f"表格 {idx + 1}:")
            
            if source_merges:
                source_has_merge += 1
                print(f"  源文档 - 找到 {len(source_merges)} 个横向合并:")
                for m in source_merges:
                    print(f"    [{m['row']},{m['col']}] 跨{m['span']}列 - '{m['text']}'")
            
            if target_merges:
                target_has_merge += 1
                print(f"  目标文档 - 找到 {len(target_merges)} 个横向合并:")
                for m in target_merges:
                    print(f"    [{m['row']},{m['col']}] 跨{m['span']}列 - '{m['text']}'")
            
            # 对比是否一致
            if len(source_merges) == len(target_merges):
                all_match = True
                for sm, tm in zip(source_merges, target_merges):
                    if sm['row'] != tm['row'] or sm['col'] != tm['col'] or sm['span'] != tm['span']:
                        all_match = False
                        break
                
                if all_match:
                    match_count += 1
                    print(f"  [OK] 横向合并结构完全匹配")
                else:
                    mismatch_count += 1
                    print(f"  [WARN] 横向合并位置或跨度不匹配")
            else:
                mismatch_count += 1
                print(f"  [ERROR] 横向合并数量不一致 (源:{len(source_merges)}, 目标:{len(target_merges)})")
    
    print(f"\n{'='*70}")
    print("总结:")
    print(f"  源文档有横向合并的表格: {source_has_merge}")
    print(f"  目标文档有横向合并的表格: {target_has_merge}")
    print(f"  完全匹配的表格: {match_count}")
    print(f"  不匹配的表格: {mismatch_count}")
    
    if mismatch_count == 0 and source_has_merge == target_has_merge:
        print("\n[SUCCESS] 所有横向合并都正确保留！")
    else:
        print(f"\n[WARNING] 发现 {mismatch_count} 个表格的横向合并不匹配")

if __name__ == "__main__":
    compare_horizontal_merge()
