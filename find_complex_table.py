# -*- coding: utf-8 -*-
"""查找包含复杂合并单元格的表格"""
from docx import Document
from docx.oxml.ns import qn

def analyze_table_structure(table, idx, doc_name):
    print("\n" + "="*80)
    print(f"{doc_name} - 表格 #{idx}")
    print("="*80)
    print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")
    print()
    
    # 显示前8行
    for i, row in enumerate(table.rows[:8]):
        print(f"行 {i+1}:")
        for j, cell in enumerate(row.cells):
            text = ''.join([p.text for p in cell.paragraphs]).strip()
            if text:
                print(f"  [{i},{j}]: '{text[:20]}'")
        
        # 检查合并属性
        for j, cell in enumerate(row.cells):
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                gridSpan = tcPr.find(qn('w:gridSpan'))
                vMerge = tcPr.find(qn('w:vMerge'))
                
                if gridSpan is not None:
                    print(f"  -> [{i},{j}]: 横向合并 gridSpan={gridSpan.get(qn('w:val'))}")
                if vMerge is not None:
                    val = vMerge.get(qn('w:val'))
                    print(f"  -> [{i},{j}]: 纵向合并 vMerge={val}")
        print()

# 搜索所有文档
keyword_groups = [
    ["DI", "DO", "AI", "AO"],
    ["消防水泵", "喷淋泵"],
]

for doc_file in ["8js1.docx", "8js2.docx", "8js3.docx"]:
    print(f"\n搜索 {doc_file}...")
    doc = Document(doc_file)
    
    for idx, table in enumerate(doc.tables, 1):
        all_text = ""
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + " "
        
        # 检查是否包含关键词
        for keywords in keyword_groups:
            if all(kw in all_text for kw in keywords):
                analyze_table_structure(table, idx, doc_file)
                print(f"\n>>> 这是你要的表格吗？")
                import sys
                sys.exit(0)

print("\n未找到匹配的表格")
