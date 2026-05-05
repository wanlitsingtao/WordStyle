# -*- coding: utf-8 -*-
"""测试 python-docx 的 columns 属性在合并单元格时的行为"""
from docx import Document
from docx.oxml.ns import qn

doc = Document("E:/LingMa/8js3.docx")

# 找到表格 #36
table_36 = doc.tables[35]  # 索引从0开始

print("="*80)
print("测试表格 #36 的列数计算")
print("="*80)
print()

# 方法1: 使用 len(table.columns)
print(f"方法1: len(table.columns) = {len(table_36.columns)}")

# 方法2: 计算第一行的最大列索引
max_col_index = 0
for cell in table_36.rows[0].cells:
    # 获取单元格的列索引
    tc = cell._element
    # 通过 tc 的兄弟节点来确定位置
    idx = list(tc.getparent()).index(tc)
    if idx > max_col_index:
        max_col_index = idx

print(f"方法2: 第一行的单元格数量 = {len(table_36.rows[0].cells)}")
print(f"方法3: 第一行的最大索引 = {max_col_index}")

# 方法3: 检查 XML 结构
print("\n分析 XML 结构:")
for i, row in enumerate(table_36.rows[:2]):
    print(f"\n行 {i+1}:")
    tc_count = 0
    for cell in row.cells:
        tc_count += 1
        text = ''.join([p.text for p in cell.paragraphs]).strip()[:10]
        
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        merge_info = ""
        if tcPr is not None:
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is not None:
                merge_info = f" (gridSpan={gridSpan.get(qn('w:val'))})"
            vMerge = tcPr.find(qn('w:vMerge'))
            if vMerge is not None:
                merge_info = f" (vMerge={vMerge.get(qn('w:val'))})"
        
        print(f"  cell[{tc_count-1}]: '{text}'{merge_info}")
    
    print(f"  该行 cell 对象数量: {tc_count}")

# 方法4: 通过 XML 直接计算网格列数
print("\n\n方法4: 通过 XML 计算实际网格列数")
for i, row in enumerate(table_36.rows[:1]):
    grid_cols = 0
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        
        # 默认占1列
        span = 1
        if tcPr is not None:
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is not None:
                span = int(gridSpan.get(qn('w:val'), 1))
        
        grid_cols += span
        print(f"  cell 占 {span} 列, 累计: {grid_cols}")
    
    print(f"\n行 {i+1} 的实际网格列数: {grid_cols}")
