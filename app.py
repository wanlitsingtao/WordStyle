# -*- coding: utf-8 -*-
"""
文档转换工具 - Web 版本 (MVP)
基于 Streamlit 快速搭建
"""
import streamlit as st
import os
import sys
import json
import threading
from datetime import datetime, timedelta
from pathlib import Path

# 添加当前目录到路径，以便导入 doc_converter
sys.path.insert(0, os.path.dirname(__file__))

from doc_converter import DocumentConverter
from task_manager import (
    create_task, update_task_status, complete_task, fail_task,
    get_user_active_task, get_user_completed_tasks, has_active_task,
    cleanup_expired_tasks, RESULTS_DIR
)

# ==================== 配置 ====================
# 计费规则：100个段落 = 0.1元
PARAGRAPH_PRICE = 0.001  # 每个段落的价格（元）
MIN_RECHARGE = 1.0  # 最低充值金额（元）
BACKEND_URL = "http://localhost:8000"  # 后端API地址（用于获取免费额度配置）

# 充值档位
RECHARGE_PACKAGES = [
    {'amount': 1, 'paragraphs': 1000, 'label': '体验版'},
    {'amount': 5, 'paragraphs': 5000, 'label': '标准版'},
    {'amount': 10, 'paragraphs': 10000, 'label': '专业版'},
    {'amount': 50, 'paragraphs': 50000, 'label': '企业版'},
    {'amount': 100, 'paragraphs': 100000, 'label': '旗舰版'},
]

ADMIN_CONTACT = "微信号：your_wechat_id"  # 管理员联系方式（请修改）

# ==================== 辅助函数 ====================

def load_user_data():
    """加载用户数据（基于浏览器会话）"""
    # 使用稳定的用户标识
    if 'user_id' not in st.session_state:
        # 生成一个稳定的用户ID（基于session，但在会话期间保持不变）
        import hashlib
        import time
        # 使用时间戳 + 随机数生成唯一ID
        unique_key = f"{time.time()}_{id(st.session_state)}"
        st.session_state.user_id = hashlib.md5(unique_key.encode()).hexdigest()[:12]
    
    user_id = st.session_state.user_id
    data_file = Path("user_data.json")
    
    if data_file.exists():
        with open(data_file, 'r', encoding='utf-8') as f:
            try:
                all_data = json.load(f)
                return all_data.get(user_id, {
                    'balance': 0.0,  # 余额（元）
                    'paragraphs_remaining': 0,  # 剩余段落数
                    'total_converted': 0,  # 累计转换文档数
                    'total_paragraphs_used': 0,  # 累计使用段落数
                    'recharge_history': [],  # 充值记录
                    'conversion_history': []  # 转换记录
                })
            except:
                return {
                    'balance': 0.0,
                    'paragraphs_remaining': 0,
                    'total_converted': 0,
                    'total_paragraphs_used': 0,
                    'recharge_history': [],
                    'conversion_history': []
                }
    else:
        return {
            'balance': 0.0,
            'paragraphs_remaining': 0,
            'total_converted': 0,
            'total_paragraphs_used': 0,
            'recharge_history': [],
            'conversion_history': []
        }

# ==================== 初始化会话状态 ====================
# 使用稳定的用户ID，防止刷新页面后重新领取免费额度
# 策略1：优先从 localStorage 持久化的用户ID（关机重启也有效）
# 策略2：从 user_data.json 中查找最近使用的用户（兜底方案）

if 'user_id' not in st.session_state:
    # 尝试从 localStorage 获取用户ID（通过隐藏组件）
    import streamlit.components.v1 as components
    
    # 创建隐藏的 HTML 组件来读取 localStorage
    # 注意：这个组件会在页面加载时执行，将用户ID写入 sessionStorage
    js_init_code = """
    <script>
        (function() {
            // 检查是否已有持久化的用户ID
            let userId = localStorage.getItem('wordstyle_user_id');
            if (!userId) {
                // 生成新的用户ID并持久化
                userId = 'user_' + Date.now().toString(36) + '_' + Math.random().toString(36).substr(2, 9);
                localStorage.setItem('wordstyle_user_id', userId);
                console.log('[WordStyle] Generated new user ID:', userId);
            } else {
                console.log('[WordStyle] Using existing user ID:', userId);
            }
            
            // 将用户ID存储到全局变量，供后续读取
            window.wordstyle_user_id = userId;
        })();
    </script>
    """
    
    try:
        # 执行初始化脚本（设置 localStorage）
        components.html(js_init_code, height=0, width=0)
    except:
        pass
    
    # 由于无法直接从 JavaScript 同步获取返回值，我们使用兜底方案
    # 从 user_data.json 中查找最近使用的用户
    data_file = Path("user_data.json")
    existing_user_id = None
    
    if data_file.exists():
        try:
            with open(data_file, 'r', encoding='utf-8') as f:
                all_data = json.load(f)
                if all_data:
                    # 优先级1：找有转换记录的用户（说明真正使用过）
                    for uid, udata in all_data.items():
                        if udata.get('total_converted', 0) > 0:
                            existing_user_id = uid
                            break
                    
                    # 优先级2：如果没有转换记录，找有充值记录的用户
                    if not existing_user_id:
                        for uid, udata in all_data.items():
                            if len(udata.get('recharge_history', [])) > 0:
                                existing_user_id = uid
                                break
                    
                    # 优先级3：如果都没有，找余额不为默认值的用户
                    if not existing_user_id:
                        for uid, udata in all_data.items():
                            if udata.get('paragraphs_remaining', 0) > 0 and udata.get('paragraphs_remaining', 0) != 10000:
                                existing_user_id = uid
                                break
        except:
            pass
    
    if existing_user_id:
        # 使用已存在的用户ID
        st.session_state.user_id = existing_user_id
        # 同步更新 localStorage（确保一致性）
        try:
            sync_js = f"""
            <script>
                localStorage.setItem('wordstyle_user_id', '{existing_user_id}');
                console.log('[WordStyle] Synced user ID to localStorage:', '{existing_user_id}');
            </script>
            """
            components.html(sync_js, height=0, width=0)
        except:
            pass
    else:
        # 生成新的用户ID
        import hashlib
        import time
        unique_key = f"{time.time()}_{id(st.session_state)}"
        new_user_id = hashlib.md5(unique_key.encode()).hexdigest()[:12]
        st.session_state.user_id = new_user_id
        
        # 保存到 localStorage
        try:
            save_js = f"""
            <script>
                localStorage.setItem('wordstyle_user_id', '{new_user_id}');
                console.log('[WordStyle] Saved new user ID to localStorage:', '{new_user_id}');
            </script>
            """
            components.html(save_js, height=0, width=0)
        except:
            pass

if 'free_paragraphs_claimed' not in st.session_state:
    # 检查该用户是否已领取过免费额度
    user_data = load_user_data()
    # 如果用户有充值记录或转换记录，说明已经使用过，不再赠送
    has_used = (
        len(user_data.get('recharge_history', [])) > 0 or
        user_data.get('total_converted', 0) > 0 or
        user_data.get('paragraphs_remaining', 0) >= 10000  # 如果余额超过默认值，说明已领取
    )
    st.session_state.free_paragraphs_claimed = has_used

def save_user_data(user_data):
    """保存用户数据"""
    user_id = st.session_state.user_id
    data_file = Path("user_data.json")
    
    all_data = {}
    if data_file.exists():
        with open(data_file, 'r', encoding='utf-8') as f:
            try:
                all_data = json.load(f)
            except:
                pass
    
    all_data[user_id] = user_data
    
    with open(data_file, 'w', encoding='utf-8') as f:
        json.dump(all_data, f, ensure_ascii=False, indent=2)

# ==================== 评论区功能 ====================

COMMENTS_FILE = Path("comments_data.json")

def load_comments():
    """加载评论数据"""
    if COMMENTS_FILE.exists():
        with open(COMMENTS_FILE, 'r', encoding='utf-8') as f:
            try:
                return json.load(f)
            except:
                return []
    return []

def save_comments(comments):
    """保存评论数据"""
    with open(COMMENTS_FILE, 'w', encoding='utf-8') as f:
        json.dump(comments, f, ensure_ascii=False, indent=2)

def add_comment(username, content, rating=5):
    """添加新评论"""
    comments = load_comments()
    
    new_comment = {
        'id': len(comments) + 1,
        'username': username or f'用户{st.session_state.user_id[:6]}',
        'content': content,
        'rating': rating,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'likes': 0,
        'user_id': st.session_state.user_id
    }
    
    comments.append(new_comment)
    save_comments(comments)
    return new_comment

def like_comment(comment_id):
    """点赞评论"""
    comments = load_comments()
    for comment in comments:
        if comment['id'] == comment_id:
            comment['likes'] += 1
            break
    save_comments(comments)

def show_comments_section():
    """显示评论区"""
    # 加载评论
    comments = load_comments()
    
    # 显示统计信息
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("总评论数", len(comments))
    with col2:
        if comments:
            avg_rating = sum(c.get('rating', 5) for c in comments) / len(comments)
            st.metric("平均评分", f"{avg_rating:.1f} ⭐")
        else:
            st.metric("平均评分", "暂无")
    with col3:
        total_likes = sum(c.get('likes', 0) for c in comments)
        st.metric("总点赞数", total_likes)
    
    st.markdown("---")
    
    # 发表评论表单
    with st.expander("✍️ 发表评论", expanded=False):
        with st.form("comment_form"):
            rating = st.slider("评分", 1, 5, 5, help="请为工具打分")
            
            content = st.text_area(
                "评论内容",
                placeholder="分享您的使用体验、建议或问题...",
                height=100,
                max_chars=500
            )
            
            col_submit, col_cancel = st.columns([1, 3])
            with col_submit:
                submit_comment = st.form_submit_button("📤 发表", type="primary")
            
            if submit_comment:
                if not content.strip():
                    st.error("❌ 请输入评论内容")
                else:
                    new_comment = add_comment(None, content, rating)  # 匿名评论
                    st.success("✅ 评论发表成功！")
                    st.session_state.comment_submitted = True
                    st.rerun()
    
    st.markdown("---")
    
    # 显示评论列表
    if not comments:
        st.info("💭 暂无评论，快来发表第一条评论吧！")
    else:
        # 按时间倒序显示
        comments_sorted = sorted(comments, key=lambda x: x['timestamp'], reverse=True)
        
        for comment in comments_sorted[:20]:  # 最多显示20条
            with st.container():
                col_header, col_like = st.columns([4, 1])
                
                with col_header:
                    # 只显示评分和时间
                    stars = "⭐" * comment.get('rating', 5)
                    st.markdown(f"{stars}")
                    st.caption(f"🕒 {comment.get('timestamp', '')}")
                
                with col_like:
                    # 点赞按钮
                    likes = comment.get('likes', 0)
                    if st.button(f"👍 {likes}", key=f"like_{comment['id']}"):
                        like_comment(comment['id'])
                        st.rerun()
                
                # 显示评论内容
                st.markdown(f"<div style='padding: 10px; background-color: #f0f2f6; border-radius: 5px; margin: 5px 0;'>{comment.get('content', '')}</div>", unsafe_allow_html=True)
                
                st.markdown("---")
        
        if len(comments) > 20:
            st.caption(f"显示最近20条评论，共 {len(comments)} 条")

def get_free_paragraphs_config():
    """从后端获取免费额度配置"""
    try:
        import requests
        response = requests.get(f"{BACKEND_URL}/api/admin/config/free-paragraphs", timeout=2)
        if response.status_code == 200:
            data = response.json()
            return data.get('config_value', 10000)
        return 10000
    except:
        return 10000  # 默认值

def claim_free_paragraphs():
    """领取免费额度（仅首次访问）"""
    # 如果已经标记为已领取，直接返回
    if st.session_state.free_paragraphs_claimed:
        return 0
    
    # 加载用户数据
    user_data = load_user_data()
    
    # 多重检查：是否已经领取过
    has_claimed = (
        len(user_data.get('recharge_history', [])) > 0 or  # 有充值记录
        user_data.get('total_converted', 0) > 0 or  # 有转换记录
        user_data.get('paragraphs_remaining', 0) >= 10000  # 余额已经达到或超过默认值
    )
    
    if has_claimed:
        st.session_state.free_paragraphs_claimed = True
        return 0
    
    # 首次访问，赠送免费额度
    free_paragraphs = get_free_paragraphs_config()
    user_data['paragraphs_remaining'] += free_paragraphs
    
    # 添加领取记录
    if 'free_quota_history' not in user_data:
        user_data['free_quota_history'] = []
    
    user_data['free_quota_history'].append({
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'amount': free_paragraphs,
        'type': 'first_login_bonus'
    })
    
    save_user_data(user_data)
    st.session_state.free_paragraphs_claimed = True
    return free_paragraphs

def count_paragraphs(docx_file):
    """统计文档段落数"""
    try:
        from docx import Document
        doc = Document(docx_file)
        return len(doc.paragraphs)
    except:
        return 0

def get_template_styles_list(template_file):
    """获取模板文档中的所有段落样式"""
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        doc = Document(template_file)
        styles = []
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                styles.append(style.name)
        return sorted(styles)
    except:
        return ["Normal"]  # 默认返回Normal样式

def analyze_source_styles_with_progress(source_files, user_id):
    """
    分析源文档样式并显示进度条（参照桌面版逻辑）
    :param source_files: 上传的文件对象列表
    :param user_id: 用户ID
    :return: {filename: [styles]} 字典，每个文件对应其样式列表
    """
    import os
    from docx import Document
    
    file_styles_map = {}  # {filename: [styles]}
    total_files = len(source_files)
    
    # 创建进度条
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for idx, source_file in enumerate(source_files, 1):
        status_text.text(f"🔍 正在分析第 {idx}/{total_files} 个文件: {source_file.name}...")
        
        # 保存临时文件
        temp_source = f"temp_source_{user_id}_{source_file.name}"
        try:
            with open(temp_source, 'wb') as f:
                f.write(source_file.getbuffer())
            
            # 读取样式
            doc = Document(temp_source)
            para_count = len(doc.paragraphs)
            styles = set()
            
            for para_idx, para in enumerate(doc.paragraphs):
                if para.style and para.style.name:
                    styles.add(para.style.name)
                
                # 每处理10个段落或最后一个段落时更新进度
                if (para_idx + 1) % 10 == 0 or para_idx == para_count - 1:
                    # 计算总体进度
                    completed_files_progress = (idx - 1) / total_files
                    current_file_progress = ((para_idx + 1) / para_count) / total_files
                    total_progress = completed_files_progress + current_file_progress
                    progress_bar.progress(min(int(total_progress * 100), 100))
            
            # 保存该文件的样式
            file_styles_map[source_file.name] = sorted(list(styles))
            
        except Exception as e:
            st.error(f"❌ 分析文件 {source_file.name} 失败: {e}")
            continue
    
    # 完成
    progress_bar.progress(100)
    status_text.text("✅ 样式分析完成！")
    
    return file_styles_map

def execute_background_conversion(task_id, source_files_info, template_path, config, user_id):
    """
    在后台线程中执行转换任务
    :param task_id: 任务ID
    :param source_files_info: 源文件信息列表 [(filename, temp_path, paragraphs), ...]
    :param template_path: 模板文件路径
    :param config: 转换配置字典
    :param user_id: 用户ID
    """
    try:
        update_task_status(task_id, 'PROCESSING')
        
        converter = DocumentConverter()
        output_files = []
        total_success_paragraphs = 0
        
        for idx, (filename, temp_source, file_paragraphs) in enumerate(source_files_info):
            # 输出文件路径
            base_name = os.path.splitext(filename)[0]
            output_filename = f"{user_id}_{task_id[:8]}_{base_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(RESULTS_DIR, output_filename)
            
            # 警告收集
            warnings_list = []
            def warning_callback(msg):
                warnings_list.append(msg)
            
            # 进度回调
            def make_progress_callback(file_idx, total_files):
                def callback(step, message):
                    # 计算总体进度 (0-100%)
                    base_progress = int((file_idx / total_files) * 100)
                    step_progress = int((step / 7) * (100 / total_files))
                    current_progress = min(base_progress + step_progress, 95)
                    update_task_status(task_id, 'PROCESSING', progress=current_progress)
                return callback
            
            # 执行转换
            success, actual_file, msg = converter.full_convert(
                source_file=temp_source,
                template_file=template_path,
                output_file=output_path,
                custom_style_map=config.get('custom_style_map', None),  # 使用配置中的样式映射
                do_mood=config['do_mood'],
                answer_text=config['answer_text'],
                answer_style=config['answer_style'],
                list_bullet=config['list_bullet'],
                do_answer_insertion=config['do_answer_insertion'],
                answer_mode=config['answer_mode'],
                progress_callback=make_progress_callback(idx, len(source_files_info)),
                warning_callback=warning_callback
            )
            
            if success:
                output_files.append(output_path)
                total_success_paragraphs += file_paragraphs
            else:
                # 转换失败，清理已生成的文件
                for of in output_files:
                    try:
                        if os.path.exists(of):
                            os.remove(of)
                    except:
                        pass
                fail_task(task_id, f"文件 {filename} 转换失败: {msg}")
                return
        
        # 所有文件转换成功
        complete_task(task_id, output_files)
        
        # 扣费（只在完全成功后扣费）
        from app import load_user_data, save_user_data, calculate_cost
        user_data = load_user_data()
        actual_cost = calculate_cost(total_success_paragraphs)
        user_data['paragraphs_remaining'] -= total_success_paragraphs
        user_data['balance'] -= actual_cost  # 同时扣除余额
        user_data['total_converted'] += len(output_files)
        user_data['total_paragraphs_used'] += total_success_paragraphs
        
        # 记录转换历史
        conversion_record = {
            'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'files': len(source_files_info),
            'success': len(output_files),
            'failed': 0,
            'paragraphs_charged': total_success_paragraphs,
            'cost': actual_cost,
            'task_id': task_id,
            'mode': 'background'
        }
        user_data['conversion_history'].append(conversion_record)
        save_user_data(user_data)
        
    except Exception as e:
        # 转换异常，清理文件
        fail_task(task_id, f"转换异常: {str(e)}")
        import traceback
        traceback.print_exc()

def calculate_cost(paragraphs):
    """计算转换费用"""
    return paragraphs * PARAGRAPH_PRICE

def recharge_user(user_data, amount, package_label):
    """用户充值"""
    # 计算获得的段落数
    paragraphs_to_add = int(amount / PARAGRAPH_PRICE)
    
    # 更新余额和段落数
    user_data['balance'] += amount
    user_data['paragraphs_remaining'] += paragraphs_to_add
    
    # 记录充值历史
    recharge_record = {
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'amount': amount,
        'paragraphs': paragraphs_to_add,
        'package': package_label
    }
    user_data['recharge_history'].append(recharge_record)
    
    # 保存
    save_user_data(user_data)
    
    return paragraphs_to_add

def count_pages(docx_file):
    """估算文档页数（基于段落数）"""
    try:
        from docx import Document
        doc = Document(docx_file)
        # 粗略估算：每50个段落约1页
        paragraphs = len(doc.paragraphs)
        estimated_pages = max(1, paragraphs // 50)
        return estimated_pages
    except:
        return 0  # 无法计算时返回0

# ==================== 页面配置 ====================
st.set_page_config(
    page_title="标书抄写神器",
    page_icon="📄",
    layout="wide",  # 使用宽屏布局
    initial_sidebar_state="expanded"
)

# ==================== 定期清理过期任务 ====================
# 每次加载页面时检查并清理过期任务（7天前的任务）
try:
    cleaned_count = cleanup_expired_tasks()
    if cleaned_count > 0:
        print(f"已清理 {cleaned_count} 个过期任务")
except Exception as e:
    print(f"清理过期任务失败: {e}")

# ==================== 主界面 ====================
st.title("📝 标书抄写神器")

# 全屏提示
st.markdown("""
<div style='background-color: #e3f2fd; padding: 10px; border-radius: 5px; margin-bottom: 10px;'>
💡 <strong>提示：</strong>按 <kbd>F11</kbd> 键可以让浏览器全屏显示，获得更好的体验
</div>
""", unsafe_allow_html=True)

st.markdown("---")

# 自定义CSS，优化页面显示（简化版，让Streamlit自动处理布局）
st.markdown("""
<style>
    /* 隐藏页脚 */
    footer {visibility: hidden;}
    
    /* 强制主要内容区域使用最大宽度 */
    .block-container {
        max-width: 100% !important;
        padding-top: 2rem !important;
        padding-bottom: 1rem !important;
    }
    
    /* 优化文件上传器大小 */
    .stFileUploader > div {
        min-height: 80px;
    }
    
    /* 增大按钮 */
    .stButton > button {
        height: 3em;
        font-size: 1.1em;
        width: 100%;
    }
    
    /* 优化指标显示 */
    [data-testid="stMetric"] {
        background-color: #f0f2f6;
        padding: 15px;
        border-radius: 5px;
    }
    
    /* 修复侧边栏隐藏后的布局问题 */
    div[data-testid="stAppViewContainer"] {
        width: 100% !important;
    }
    
    /* 确保主内容区域的父容器正确响应侧边栏变化 */
    section[data-testid="stSidebar"] + div {
        flex-grow: 1 !important;
        width: auto !important;
    }
</style>

<script>
// 监听侧边栏按钮点击并强制重新布局
setTimeout(function() {
    // 查找侧边栏切换按钮
    const toggleButtons = document.querySelectorAll('button[title*="sidebar"], button[aria-label*="sidebar"], [data-testid="stSidebarCollapsedControl"]');
    
    toggleButtons.forEach(function(btn) {
        btn.addEventListener('click', function() {
            // 延迟执行以确保DOM已更新
            setTimeout(function() {
                // 触发窗口resize事件
                window.dispatchEvent(new Event('resize'));
                
                // 强制重新计算布局
                const mainContainer = document.querySelector('.main');
                if (mainContainer) {
                    mainContainer.style.display = 'none';
                    setTimeout(function() {
                        mainContainer.style.display = '';
                    }, 10);
                }
            }, 300);
        });
    });
}, 2000);
</script>
""", unsafe_allow_html=True)

# 侧边栏：用户信息
with st.sidebar:
    st.header("👤 用户信息")
    
    # 首次访问，自动领取免费额度
    if not st.session_state.free_paragraphs_claimed:
        free_paragraphs = claim_free_paragraphs()
        if free_paragraphs > 0:
            st.toast(f"🎉 欢迎！已赠送您 {free_paragraphs:,} 段免费额度", icon="🎁")
    
    # 加载用户数据
    user_data = load_user_data()
    
    # 显示余额和段落数
    st.metric("账户余额", f"¥{user_data['balance']:.2f}")
    st.metric("剩余段落数", f"{user_data['paragraphs_remaining']:,}")
    st.metric("累计转换文档", user_data['total_converted'])
    
    st.markdown("---")
    
    # ==================== 暂时隐藏充值功能 ====================
    # TODO: 等支付功能成熟后再启用
    if False:  # 临时禁用充值功能
        # 微信扫码充值
        st.markdown("### 💳 扫码充值")
        
        # 充值档位选择
        recharge_options = [f"{pkg['label']} - ¥{pkg['amount']} ({pkg['paragraphs']:,}段)" for pkg in RECHARGE_PACKAGES]
        selected_package = st.selectbox("选择充值档位", recharge_options, label_visibility="collapsed")
        
        if st.button("生成支付二维码", type="primary", use_container_width=True):
            # 解析选择的套餐
            for pkg in RECHARGE_PACKAGES:
                if f"{pkg['label']} - ¥{pkg['amount']}" in selected_package:
                    amount = pkg['amount']
                    paragraphs = pkg['paragraphs']
                    package_label = pkg['label']
                    break
            
            # 调用后端API创建订单
            try:
                import requests
                response = requests.post(
                    f"{BACKEND_URL}/api/test-payment/create",
                    json={
                        'user_id': st.session_state.user_id,
                        'amount': amount,
                        'paragraphs': paragraphs,
                        'package_label': package_label
                    },
                    timeout=5
                )
                
                if response.status_code == 200:
                    order_data = response.json()
                    order_id = order_data['order_id']
                    
                    st.success(f"✅ 订单已创建！请扫描下方二维码转账 ¥{amount}")
                    st.info("💡 扫码支付后，请点击下方的'✅ 我已支付'按钮")
                    
                    # 显示您的个人收款码
                    qr_code_path = "personal_qr_code.png"  # 您的收款码图片文件
                    
                    if Path(qr_code_path).exists():
                        # 如果存在收款码图片，显示它
                        st.image(qr_code_path, caption=f"微信扫码支付 ¥{amount}\n\n订单号: {order_id}", width=280)
                        
                        # 添加使用说明
                        with st.expander("📖 如何充值？", expanded=True):
                            st.markdown(f"""
**充值步骤：**
1. 打开微信，扫描上方二维码
2. 输入金额 **¥{amount}**
3. 完成支付
4. 点击下方“✅ 我已支付”按钮
5. 系统将自动为您充值 **{paragraphs:,}** 段

**订单信息：**
- 订单号：`{order_id}`
- 用户ID：`{st.session_state.user_id[:8]}...`
- 充值金额：¥{amount}
- 获得段落：{paragraphs:,} 段
                        """)
                    else:
                        # 否则显示配置提示
                        st.warning("⚠️ **管理员尚未配置收款码**")
                        st.info("💡 请将微信收款码截图保存为 `personal_qr_code.png` 文件放在项目根目录")
                        
                        # 提供上传功能（仅演示）
                        uploaded_file = st.file_uploader("上传收款码图片", type=['png', 'jpg', 'jpeg'])
                        if uploaded_file:
                            with open(qr_code_path, 'wb') as f:
                                f.write(uploaded_file.read())
                            st.success("✅ 收款码已保存！刷新页面即可使用")
                            st.rerun()
                    
                    # 添加“我已支付”按钮
                    st.markdown("---")
                                    
                    if Path(qr_code_path).exists():
                        st.caption("💡 **重要提示**：请确认您已完成支付后再点击下方按钮")
                                        
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            if st.button("✅ 我已支付（确认充值）", type="primary", use_container_width=True, key=f"confirm_pay_{order_id}"):
                                with st.spinner("正在处理充值..."):
                                    time.sleep(0.5)  # 模拟处理时间
                                                    
                                    # 调用后端API完成充值
                                    pay_response = requests.post(
                                        f"{BACKEND_URL}/api/test-payment/simulate-payment",
                                        json={
                                            'order_id': order_id,
                                            'user_id': st.session_state.user_id
                                        },
                                        timeout=5
                                    )
                                                    
                                    if pay_response.status_code == 200:
                                        result = pay_response.json()
                                        st.balloons()  # 🎈 彩带庆祝
                                        st.success(f"✅ {result['message']}")
                                        st.info(f"💰 已充值 {paragraphs:,} 段，当前余额: ¥{result.get('new_balance', 0):.2f}")
                                                        
                                        # 强制刷新页面以显示新余额
                                        time.sleep(1.5)
                                        st.rerun()
                                    else:
                                        st.error("❌ 充值失败，请联系管理员")
                                        
                        with col2:
                            if st.button("🔄 重新生成", use_container_width=True):
                                st.rerun()
                    
                    with col2:
                        if st.button("❌ 取消", use_container_width=True, key=f"cancel_{order_id}"):
                            st.info("订单已取消")
                            time.sleep(1)
                            st.rerun()
                    
                    # 轮询检查支付状态（每3秒检查一次）
                    status_placeholder = st.empty()
                    for _ in range(20):  # 最多检查60秒
                        time.sleep(3)
                        check_response = requests.get(
                            f"{BACKEND_URL}/api/test-payment/check-status/{order_id}",
                            timeout=3
                        )
                        
                        if check_response.status_code == 200:
                            status = check_response.json().get('status')
                            if status == 'PAID':
                                status_placeholder.success("✅ 支付成功！页面将自动刷新...")
                                time.sleep(1)
                                st.rerun()
                                break
                            elif status == 'PENDING':
                                status_placeholder.info("⏳ 等待您确认支付...")
                else:
                    st.error(f"❌ 创建订单失败：{response.text}")
            
            except Exception as e:
                st.error(f"❌ 网络错误：{str(e)}")
                st.info("💡 请确保后端服务正在运行（http://localhost:8000）")
    
    st.markdown("---")
    
    # 管理员联系方式
    st.caption(f"📞 如需帮助，请联系：{ADMIN_CONTACT}")
    
    st.markdown("---")
    
    # 需求提交入口
    if st.button("💡 提交需求/反馈", use_container_width=True):
        st.session_state.show_feedback_form = True
        st.rerun()
    
    # 管理后台入口（隐藏链接，通过URL访问）
    # st.markdown("[🔧 管理后台](/?page=admin)")
    
    st.markdown("---")
    st.caption("© 2026 文档转换工具 | 按量付费")

# ==================== 主功能区 ====================

# 需求提交表单对话框
if st.session_state.get('show_feedback_form'):
    with st.form("feedback_form"):
        st.header("💡 提交需求或反馈")
        st.markdown("我们非常重视您的意见，请告诉我们您的想法！")
        
        # 反馈类型
        feedback_type = st.selectbox(
            "反馈类型",
            ["功能建议", "Bug报告", "使用问题", "其他"],
            help="请选择反馈的类型"
        )
        
        # 标题
        feedback_title = st.text_input(
            "标题",
            placeholder="简要描述您的需求或问题",
            help="请用一句话概括"
        )
        
        # 详细描述
        feedback_description = st.text_area(
            "详细描述",
            placeholder="请详细描述您的需求、问题或建议...\n\n例如：\n- 我希望增加XX功能\n- 我遇到了XX问题\n- 我觉得XX可以改进",
            height=150,
            help="越详细越好，帮助我们更好地理解您的需求"
        )
        
        # 联系方式（可选）
        feedback_contact = st.text_input(
            "联系方式（可选）",
            placeholder="微信/邮箱/电话",
            help="如果需要我们回复您，请留下联系方式"
        )
        
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            submit_feedback = st.form_submit_button("✅ 提交", type="primary", use_container_width=True)
        with col2:
            cancel_feedback = st.form_submit_button("❌ 取消", use_container_width=True)
    
    if cancel_feedback:
        st.session_state.show_feedback_form = False
        st.rerun()
    
    if submit_feedback:
        if not feedback_title or not feedback_description:
            st.error("❌ 请填写标题和详细描述")
        else:
            try:
                import requests
                
                # 映射反馈类型
                type_map = {
                    "功能建议": "feature",
                    "Bug报告": "bug",
                    "使用问题": "question",
                    "其他": "other"
                }
                
                response = requests.post(
                    f"{BACKEND_URL}/api/feedback/submit",
                    json={
                        'user_id': st.session_state.user_id,
                        'feedback_type': type_map.get(feedback_type, 'other'),
                        'title': feedback_title,
                        'description': feedback_description,
                        'contact': feedback_contact
                    },
                    timeout=5
                )
                
                if response.status_code == 200:
                    result = response.json()
                    st.success(f"✅ {result['message']}")
                    st.info(f"📝 反馈ID: {result['feedback_id']}")
                    st.session_state.show_feedback_form = False
                    
                    # 2秒后刷新
                    import time
                    time.sleep(2)
                    st.rerun()
                else:
                    st.error(f"❌ 提交失败：{response.text}")
            except Exception as e:
                st.error(f"❌ 网络错误：{str(e)}")
                st.info("💡 请确保后端服务正在运行")
    
    st.markdown("---")

# 文件上传区
col1, col2 = st.columns([1.2, 1.2])  # 等宽两列

with col1:
    st.subheader("📤 上传源文档")
    source_files = st.file_uploader(
        "选择要转换的 Word 文档（可多选）",
        type=['docx'],
        help="支持 .docx 格式，可同时选择多个文件",
        accept_multiple_files=True,
        key="source_uploader"
    )
    
    if source_files:
        # 保存到 session_state，供对话框使用
        st.session_state.current_source_files = source_files
        
        # 显示上传的文件列表
        st.success(f"✅ 已上传 {len(source_files)} 个文件")
        
        # 检查是否需要重新分析（文件变化或尚未分析）
        need_analyze = False
        current_file_names = [sf.name for sf in source_files]
        analyzed_file_names = list(st.session_state.get('file_styles_map', {}).keys())
        
        if not analyzed_file_names or set(current_file_names) != set(analyzed_file_names):
            need_analyze = True
        
        if need_analyze:
            # 立即分析源文档样式（参照桌面版逻辑，带进度条）
            file_styles_map = analyze_source_styles_with_progress(source_files, st.session_state.user_id)
            st.session_state.file_styles_map = file_styles_map
            
            # 合并所有文件的样式用于显示
            all_styles = set()
            for styles in file_styles_map.values():
                all_styles.update(styles)
            all_styles = sorted(list(all_styles))
            st.session_state.source_styles = all_styles
            
            st.info(f"📋 共检测到 {len(all_styles)} 种样式: {', '.join(all_styles[:5])}{'...' if len(all_styles) > 5 else ''}")
        else:
            # 使用已缓存的样式
            file_styles_map = st.session_state.file_styles_map
            all_styles = st.session_state.source_styles
            st.info(f"📋 已缓存 {len(all_styles)} 种样式")
        
        # 统计总段落数和费用
        total_paragraphs = 0
        file_info = []
        
        for sf in source_files:
            # 保存临时文件
            temp_source = f"temp_source_{st.session_state.user_id}_{sf.name}"
            with open(temp_source, 'wb') as f:
                f.write(sf.getbuffer())
            
            # 统计段落数
            paragraphs = count_paragraphs(temp_source)
            total_paragraphs += paragraphs
            file_info.append((sf.name, paragraphs))
        
        cost = calculate_cost(total_paragraphs)
        
        # 显示文件详情
        with st.expander("📄 查看文件详情"):
            for fname, fpara in file_info:
                st.text(f"• {fname}: {fpara:,} 个段落")
        
        st.info(f"📊 总段落数: {total_paragraphs:,} 个 | 预计费用: ¥{cost:.2f}")
        
        if total_paragraphs > user_data['paragraphs_remaining']:
            st.error(f"❌ 余额不足！需要 {total_paragraphs:,} 个段落，剩余 {user_data['paragraphs_remaining']:,} 个")
            st.warning("请先充值")

with col2:
    st.subheader("📋 上传模板文档")
    template_file = st.file_uploader(
        "选择模板文档",
        type=['docx'],
        help="用于定义目标样式的 Word 文档",
        key="template_uploader"
    )
    
    if template_file:
        # 保存临时文件
        temp_template = f"temp_template_{st.session_state.user_id}.docx"
        with open(temp_template, 'wb') as f:
            f.write(template_file.getbuffer())
        
        # 保存到 session_state，供对话框使用
        st.session_state.current_temp_template = temp_template
        
        st.success(f"✅ 已上传: {template_file.name}")
        
        # 检查是否需要重新分析模板样式
        if 'template_styles' not in st.session_state or st.session_state.get('last_template_name') != template_file.name:
            # 分析模板样式（带进度条）
            with st.spinner("🔍 正在分析模板样式..."):
                template_styles = get_template_styles_list(temp_template)
                st.session_state.template_styles = template_styles
                st.session_state.last_template_name = template_file.name
                st.info(f"📋 检测到 {len(template_styles)} 种段落样式")
        else:
            # 使用已缓存的样式
            template_styles = st.session_state.template_styles
            st.info(f"📋 已缓存 {len(template_styles)} 种段落样式")

# 转换配置
st.markdown("---")
st.subheader("⚙️ 转换配置")

# 第一行：四个选项横向等距分布（中线对齐）
# 使用CSS实现控件垂直居中对齐
st.markdown("""
<style>
    /* 让所有列内的元素容器垂直居中 */
    [data-testid="column"] > div {
        display: flex;
        align-items: center;
        justify-content: center;
        height: 100%;
        min-height: 40px;
    }
</style>
""", unsafe_allow_html=True)

col1, col2, col3, col4 = st.columns(4)

with col1:
    if st.button("📊 样式映射", key="open_style_mapping_btn", use_container_width=True, help="如果不采用系统给的默认配置，可自定义样式映射"):
        st.session_state.show_style_mapping_dialog = True

with col2:
    do_mood = st.checkbox("祈使语气转换", value=True, help="将文档中的祈使语气转换为投标人语气", key="mood_checkbox")

with col3:
    do_answer = st.checkbox("插入应答句", value=True, help="在标题后插入应答句", key="answer_checkbox")

with col4:
    list_bullet = st.text_input("列表符号", value="•", help="列表段落的符号", key="bullet_input")

# 第二行：应答句详细配置（仅当勾选“插入应答句”时显示）
if do_answer:
    st.markdown("**📝 应答句配置**")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        answer_text = st.text_input(
            "应答句文本",
            value="应答：本投标人理解并满足要求。",
            help="插入的应答句内容",
            key="answer_text_input"
        )
    
    with col2:
        # 获取模板样式列表
        template_styles = st.session_state.get('template_styles', ["Normal"])
        
        answer_style = st.selectbox(
            "应答句样式",
            options=template_styles,
            index=0 if "Normal" in template_styles else 0,
            help="应答句的段落样式",
            key="answer_style_select"
        )
    
    with col3:
        answer_mode_options = {
            'before_heading': '章节前插入',
            'after_heading': '章节后插入',
            'copy_chapter': '章节招标原文+应答句+招标原文副本',
            'before_paragraph': '逐段前应答',
            'after_paragraph': '逐段后应答'
        }
        answer_mode = st.selectbox(
            "插入模式",
            options=list(answer_mode_options.keys()),
            format_func=lambda x: answer_mode_options[x],
            index=0,
            help="应答句的插入位置模式",
            key="answer_mode_select"
        )
else:
    # 默认值（不插入应答句时使用）
    answer_text = "应答：本投标人理解并满足要求。"
    answer_style = "Normal"
    answer_mode = 'before_heading'

# 开始转换按钮
st.markdown("---")

# 检查是否有进行中的任务
active_task = has_active_task(st.session_state.user_id)

if active_task:
    st.warning(f"⚠️ **您有一个进行中的后台任务**\n\n文件名：{active_task['filename']}\n状态：{active_task['status']}\n\n建议等待当前任务完成后再提交新任务。")
    
    # 显示当前任务进度
    if active_task['status'] == 'PROCESSING':
        st.progress(active_task['progress'] / 100.0)
        st.text(f"转换进度：{active_task['progress']}%")
    
    st.info("💡 您可以在下方的「转换历史」中查看任务状态和下载完成的文件")
else:
    # 检查是否正在前台转换中
    is_converting = st.session_state.get('is_converting', False)
    
    if is_converting:
        # 如果正在转换，显示提示信息
        st.warning("⏳ **正在进行前台转换，请稍候...**\n\n转换期间无法进行其他操作，请耐心等待转换完成。")
        st.info("💡 转换完成后将自动恢复操作权限")
        # 不立即重置标志，让转换流程自然结束
    else:
        # 正常状态，显示开始转换按钮
        # 如果有后台任务，禁用按钮
        button_disabled = has_active_task(st.session_state.user_id)
        
        if st.button("🚀 开始转换", type="primary", use_container_width=True, disabled=button_disabled):
            # 验证输入
            if not source_files or not template_file:
                st.error("❌ 请上传源文档和模板文档")
            elif 'temp_template' not in locals():
                st.error("❌ 文件上传失败，请重试")
            else:
                # 检查是否有进行中的任务（防止重复提交）
                if has_active_task(st.session_state.user_id):
                    active_task = get_user_active_task(st.session_state.user_id)
                    st.error(f"❌ 您已有进行中的任务：{active_task['filename']}")
                    st.info("请等待当前任务完成后再提交新任务")
                    st.stop()
                
                # 设置转换标志，禁用后续操作
                st.session_state.is_converting = True
            
            # 统计总段落数并检查余额
            total_paragraphs = sum(count_paragraphs(f"temp_source_{st.session_state.user_id}_{sf.name}") for sf in source_files)
            cost = calculate_cost(total_paragraphs)
            
            if total_paragraphs > user_data['paragraphs_remaining']:
                st.error(f"❌ 余额不足！需要 {total_paragraphs:,} 个段落（¥{cost:.2f}），剩余 {user_data['paragraphs_remaining']:,} 个")
                st.info("💡 请在左侧充值中心充值")
                st.stop()
            
            # 再次检查是否有进行中的任务（防止并发提交）
            if has_active_task(st.session_state.user_id):
                st.error("❌ 您已有进行中的任务，请等待完成后再提交")
                st.stop()
            
            # 准备文件信息
            source_files_info = []
            for sf in source_files:
                temp_source = f"temp_source_{st.session_state.user_id}_{sf.name}"
                file_paragraphs = count_paragraphs(temp_source)
                source_files_info.append((sf.name, temp_source, file_paragraphs))
            
            # 配置字典
            config = {
                'do_mood': do_mood,
                'answer_text': answer_text,
                'answer_style': answer_style,
                'list_bullet': list_bullet if list_bullet else "•",
                'do_answer_insertion': do_answer,
                'answer_mode': answer_mode,
                'custom_style_map': st.session_state.get('style_mapping', None)  # 用户配置的样式映射
            }
                        
            # ========== 前台转换模式 ==========
            # 显示进度
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # 添加“转为后台”按钮（使用session_state标记）
            if 'switch_to_background' not in st.session_state:
                st.session_state.switch_to_background = False
            
            try:
                status_text.text("⏳ 正在初始化转换器...")
                progress_bar.progress(5)
                
                # 显示“转为后台”按钮
                cancel_col1, cancel_col2 = st.columns([4, 1])
                with cancel_col2:
                    if st.button("⏸️ 转为后台", key="switch_to_bg_btn"):
                        # 显示确认对话框
                        st.warning("📝 **文档正在赶来的路上！**\n\n是否确认转为后台转换？\n- 点击**是**：转为后台转换，转换完成后从转换历史记录里下载\n- 点击**否**：继续等待前台转换完成")
                        
                        confirm_col1, confirm_col2 = st.columns(2)
                        with confirm_col1:
                            if st.button("✅ 是，转为后台", key="confirm_switch_yes", type="primary"):
                                st.session_state.switch_to_background = True
                                st.rerun()
                        with confirm_col2:
                            if st.button("❌ 否，继续等待", key="confirm_switch_no"):
                                # 不设置标记，继续前台转换
                                pass
                
                # 如果用户点击了“转为后台”，中断当前转换并创建后台任务
                if st.session_state.switch_to_background:
                    # 重置标记
                    st.session_state.switch_to_background = False
                    
                    # 创建后台任务
                    filename_display = ", ".join([sf.name for sf in source_files[:3]])
                    if len(source_files) > 3:
                        filename_display += f" 等{len(source_files)}个文件"
                    
                    task_id = create_task(
                        user_id=st.session_state.user_id,
                        filename=filename_display,
                        file_count=len(source_files),
                        paragraphs=total_paragraphs,
                        cost=cost
                    )
                    
                    # 启动后台线程
                    thread = threading.Thread(
                        target=execute_background_conversion,
                        args=(task_id, source_files_info, temp_template, config, st.session_state.user_id),
                        daemon=True
                    )
                    thread.start()
                    
                    st.success(f"✅ 已转为后台转换！任务ID: {task_id[:8]}")
                    st.info("您可以关闭此页面，稍后在「转换历史」中查看结果和下载文件")
                    st.warning("⚠️ 注意：每个用户只能有1个进行中的任务")
                    
                    # 重置转换标志
                    st.session_state.is_converting = False
                    
                    # 重要：立即停止，不再执行后续的前台转换代码
                    st.stop()
                
                # 创建转换器
                converter = DocumentConverter()
                progress_bar.progress(10)
                
                # 处理每个文件
                output_files = []
                success_count = 0
                fail_count = 0
                total_success_paragraphs = 0  # 成功转换的段落数
                
                for idx, source_file_obj in enumerate(source_files):
                    # 输出文件路径
                    base_name = os.path.splitext(source_file_obj.name)[0]
                    output_file = f"result_{base_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    temp_source = f"temp_source_{st.session_state.user_id}_{source_file_obj.name}"
                    
                    file_paragraphs = count_paragraphs(temp_source)
                    status_text.text(f"⏳ 正在转换第 {idx+1}/{len(source_files)} 个文件: {source_file_obj.name} ({file_paragraphs:,} 段落)")
                    
                    # 警告收集
                    warnings_list = []
                    def warning_callback(msg):
                        warnings_list.append(msg)
                    
                    # 进度回调函数 - 实时更新进度条
                    def make_progress_callback(file_idx, total_files):
                        def callback(step, message):
                            # 计算总体进度 (10% - 80%)
                            base_progress = 10 + int((file_idx / total_files) * 70)
                            step_progress = int((step / 7) * (70 / total_files))
                            current_progress = min(base_progress + step_progress, 80)
                            progress_bar.progress(current_progress)
                            status_text.text(f"⏳ {message}")
                        return callback
                    
                    # 执行转换
                    success, actual_file, msg = converter.full_convert(
                        source_file=temp_source,
                        template_file=temp_template,
                        output_file=output_file,
                        custom_style_map=st.session_state.get('style_mapping', None),  # 使用用户配置的样式映射
                        do_mood=do_mood,
                        answer_text=answer_text,
                        answer_style=answer_style,
                        list_bullet=list_bullet if list_bullet else "•",
                        do_answer_insertion=do_answer,
                        answer_mode=answer_mode,
                        progress_callback=make_progress_callback(idx, len(source_files)),
                        warning_callback=warning_callback
                    )
                    
                    if success:
                        output_files.append(actual_file)
                        success_count += 1
                        total_success_paragraphs += file_paragraphs
                        st.success(f"✅ {source_file_obj.name} 转换成功")
                        
                        # 显示警告（如果有）
                        if warnings_list:
                            with st.expander(f"⚠️ {source_file_obj.name} 的警告信息"):
                                for warning in warnings_list:
                                    st.warning(warning)
                    else:
                        st.error(f"❌ {source_file_obj.name} 转换失败: {msg}")
                        fail_count += 1
                
                progress_bar.progress(90)
                
                if success_count > 0:
                    status_text.text("✅ 转换完成！")
                    progress_bar.progress(100)
                    
                    # 谨慎扣费：只扣除成功转换文件的段落数
                    actual_cost = calculate_cost(total_success_paragraphs)
                    user_data['paragraphs_remaining'] -= total_success_paragraphs
                    user_data['balance'] -= actual_cost  # 同时扣除余额
                    user_data['total_converted'] += success_count
                    user_data['total_paragraphs_used'] += total_success_paragraphs
                    
                    # 记录转换历史
                    conversion_record = {
                        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'files': len(source_files),
                        'success': success_count,
                        'failed': fail_count,
                        'paragraphs_charged': total_success_paragraphs,  # 实际扣费的段落数
                        'cost': actual_cost,
                        'mode': 'foreground'
                    }
                    user_data['conversion_history'].append(conversion_record)
                    
                    # 保存用户数据
                    save_user_data(user_data)
                    
                    # 重置转换标志
                    st.session_state.is_converting = False
                    
                    # 显示结果
                    st.success(f"🎉 转换完成！成功: {success_count} 个，失败: {fail_count} 个")
                    if fail_count > 0:
                        st.warning(f"⚠️ 有 {fail_count} 个文件转换失败，未收取费用")
                    st.info(f"消耗 {total_success_paragraphs:,} 个段落（¥{actual_cost:.2f}）")
                    
                    # 提供下载
                    for output_file in output_files:
                        if os.path.exists(output_file):
                            with open(output_file, 'rb') as f:
                                file_name = os.path.basename(output_file)
                                st.download_button(
                                    label=f"📥 下载: {file_name}",
                                    data=f.read(),
                                    file_name=file_name,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
        
            except Exception as e:
                # 重置转换标志
                st.session_state.is_converting = False
                
                st.error(f"发生错误: {str(e)}")
                import traceback
                st.code(traceback.format_exc())

# ==================== 转换历史 ====================
st.markdown("---")
st.subheader("📋 我的转换历史")

# 显示保留期说明
st.info("ℹ️ **提示：** 转换完成的文件将保留 7 天，过期后会自动清理。请及时下载您需要的文件。")

# 获取用户的历史任务
completed_tasks = get_user_completed_tasks(st.session_state.user_id, limit=10)

if completed_tasks:
    for task in completed_tasks:
        with st.expander(f"{task['filename']} - {task['created_at'][:19]}"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.text(f"状态: {task['status']}")
                st.text(f"文件数: {task['file_count']}")
            
            with col2:
                if task['paragraphs']:
                    st.text(f"段落数: {task['paragraphs']:,}")
                if task['cost']:
                    st.text(f"费用: ¥{task['cost']:.2f}")
            
            with col3:
                st.text(f"创建时间: {task['created_at'][:19]}")
                if task['completed_at']:
                    st.text(f"完成时间: {task['completed_at'][:19]}")
            
            # 如果任务成功，提供下载按钮
            if task['status'] == 'COMPLETED' and task['output_files']:
                st.success("✅ 转换成功！")
                
                for output_file in task['output_files']:
                    if os.path.exists(output_file):
                        with open(output_file, 'rb') as f:
                            file_name = os.path.basename(output_file)
                            st.download_button(
                                label=f"📥 下载: {file_name}",
                                data=f.read(),
                                file_name=file_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key=f"download_{task['task_id']}_{file_name}"
                            )
            elif task['status'] == 'FAILED':
                st.error(f"❌ 转换失败: {task['error_message']}")
else:
    st.info("暂无转换历史记录")

# ==================== 使用说明 ====================
with st.expander("📖 使用说明"):
    st.markdown("""
    ### 如何使用：
    
    1. **充值**：在左侧充值中心选择档位并充值
    2. **上传源文档**：选择需要转换样式的 Word 文档
    3. **上传模板**：选择定义了目标样式的模板文档
    4. **配置选项**（可选）：
       - 语气转换：将祈使语气转换为投标人语气
       - 插入应答句：选择你需要插入应答句的位置
       - 列表符号：自定义列表段落的符号
    5. **点击开始转换**：系统会自动计算费用并扣除
    6. **下载结果**：下载转换后的文档
    
    ### 计费规则：
    - **100个段落 = ¥0.1元**
    - 按实际段落数计费，用多少扣多少
    - 余额永久有效，不会过期
    
    ### 充值档位：
    - 体验版：¥1 = 1,000段落
    - 标准版：¥5 = 5,000段落
    - 专业版：¥10 = 10,000段落
    - 企业版：¥50 = 50,000段落
    - 旗舰版：¥100 = 100,000段落
    
    ### 常见问题：
    
    **Q: 如何查看我的余额？**  
    A: 在左侧边栏可以看到账户余额和剩余段落数
    
    **Q: 转换失败会扣费吗？**  
    A: 不会！只有转换成功才会扣费
    
    **Q: 余额会过期吗？**  
    A: 不会，余额永久有效
    
    **Q: 如何获得更多优惠？**  
    A: 充值越多越划算，建议直接充值大额套餐
    
    **Q: 我的数据会保存吗？**  
    A: 仅保存转换次数统计和充值记录，文档内容不会上传到服务器
    """)

# ==================== 评论区 ====================
st.markdown("---")
st.subheader("💬 用户评论")
show_comments_section()

# ==================== 样式映射对话框 ====================
@st.dialog("📊 样式映射配置", width="large")
def show_style_mapping_dialog():
    """显示样式映射配置对话框（使用Streamlit原生dialog）"""
    st.markdown("**请为源文档中的每个样式选择对应的模板样式：**")
    st.markdown("_（未配置的样式将使用系统默认映射规则）_")
    
    # 从 session_state 获取已分析的样式
    file_styles_map = st.session_state.get('file_styles_map', {})
    template_styles = st.session_state.get('template_styles', [])
    source_files = st.session_state.get('current_source_files', None)
    
    if not file_styles_map or not source_files:
        st.warning("⚠️ 请先上传源文档并等待样式分析完成")
        return
    
    if not template_styles:
        st.warning("⚠️ 请先上传模板文档")
        return
    
    # 初始化或加载样式映射（按文件分别存储）
    if 'file_style_mappings' not in st.session_state:
        st.session_state.file_style_mappings = {}
    
    # 如果有多个文件，先选择要配置的文件
    selected_file = None
    if len(source_files) > 1:
        file_options = [sf.name for sf in source_files]
        selected_file_name = st.selectbox("选择要配置的文件", file_options, key="style_mapping_file_selector")
        selected_file = next(sf for sf in source_files if sf.name == selected_file_name)
    else:
        selected_file = source_files[0]
    
    # 获取该文件的样式列表
    source_styles = file_styles_map.get(selected_file.name, [])
    
    if not source_styles:
        st.warning(f"⚠️ 文件 {selected_file.name} 中没有检测到段落样式")
        return
    
    # 获取该文件的当前映射配置
    if selected_file.name not in st.session_state.file_style_mappings:
        st.session_state.file_style_mappings[selected_file.name] = {}
    
    current_mapping = st.session_state.file_style_mappings[selected_file.name]
    
    # 为每个源样式创建映射行（参照桌面版逻辑）
    updated_mapping = {}
    for source_style in source_styles:
        col1, col2, col3 = st.columns([2.5, 2.5, 1])
        
        with col1:
            st.text(source_style)
        
        with col2:
            # 默认值：如果当前有配置则使用，否则如果模板中有同名样式则使用，否则用Normal
            if source_style in current_mapping:
                default_value = current_mapping[source_style]
            elif source_style in template_styles:
                default_value = source_style
            else:
                default_value = "Normal"
            
            selected = st.selectbox(
                "→",
                options=template_styles,
                index=template_styles.index(default_value) if default_value in template_styles else 0,
                key=f"mapping_{selected_file.name}_{source_style}",
                label_visibility="collapsed"
            )
            updated_mapping[source_style] = selected
        
        with col3:
            hint = "✓ 已配置" if source_style in current_mapping else "○ 使用默认"
            color = "green" if source_style in current_mapping else "gray"
            st.markdown(f"<span style='color:{color};font-size:0.9em;'>{hint}</span>", unsafe_allow_html=True)
    
    # 保存更新后的映射
    st.session_state.file_style_mappings[selected_file.name] = updated_mapping
    
    # 操作按钮
    st.markdown("---")
    btn_col1, btn_col2, btn_col3 = st.columns(3)
    
    with btn_col1:
        if st.button("✅ 确定", key="confirm_mapping_btn", type="primary", use_container_width=True):
            st.success("✅ 样式映射已保存！")
            st.rerun()
    
    with btn_col2:
        if st.button("🔄 恢复默认", key="reset_mapping_btn", use_container_width=True):
            st.session_state.file_style_mappings[selected_file.name] = {}
            st.rerun()
    
    with btn_col3:
        if st.button("❌ 取消", key="cancel_mapping_btn", use_container_width=True):
            st.rerun()

# 在转换配置区调用对话框
if st.session_state.get('show_style_mapping_dialog', False):
    show_style_mapping_dialog()
    st.session_state.show_style_mapping_dialog = False  # 重置标记

# ==================== 页脚 ====================
st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: gray;'>Powered by Streamlit | MVP Version</div>",
    unsafe_allow_html=True
)
