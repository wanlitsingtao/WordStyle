# -*- coding: utf-8 -*-
"""测试第8章转换问题"""
from docx import Document
from docx.oxml.ns import qn
from doc_converter import DocumentConverter

def test_conversion(template_file, output_file):
    print(f"\n{'='*60}")
    print(f"测试模板: {template_file}")
    print(f"输出文件: {output_file}")
    print(f"{'='*60}\n")

    converter = DocumentConverter()

    # 执行转换
    success, msg = converter.convert_styles(
        source_file='8js3.docx',
        template_file=template_file,
        output_file=output_file
    )

    print(f"转换结果: {'成功' if success else '失败'}")
    print(f"消息: {msg}\n")

    if not success:
        return

    # 检查输出文件中第8章相关段落
    output_doc = Document(output_file)

    print(f"=== {output_file} 中第8章相关段落 ===")
    found_chapter8 = False
    for i, para in enumerate(output_doc.paragraphs):
        if '8.' in para.text and ('接口' in para.text or '系统' in para.text):
            outline_level = converter.get_outline_level(para)
            print(f"段落{i}: [{para.style.name}] 大纲级别:{outline_level} 文本:{para.text[:50]}")
            if '8.牵引供电' in para.text:
                found_chapter8 = True

    if not found_chapter8:
        print("未找到第8章主标题！")

    # 统计Heading样式使用情况
    heading_counts = {}
    for para in output_doc.paragraphs:
        style_name = para.style.name
        if 'Heading' in style_name:
            heading_counts[style_name] = heading_counts.get(style_name, 0) + 1

    print(f"\n=== Heading样式使用统计 ===")
    for style, count in sorted(heading_counts.items()):
        print(f"{style}: {count}个段落")

if __name__ == '__main__':
    # 测试两个模板
    test_conversion('jn.docx', 'output_jn.docx')
    test_conversion('mb.docx', 'output_mb.docx')

    print("\n\n对比分析完成！")
