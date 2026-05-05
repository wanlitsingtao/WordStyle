#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""测试读取8js2.docx文件的样式"""

from docx import Document
import os

# 检查文件是否存在
test_file = "8js2.docx"
if not os.path.exists(test_file):
    print(f"文件不存在: {test_file}")
    exit(1)

print(f"文件大小: {os.path.getsize(test_file)} bytes")
print("开始读取文档...")

try:
    doc = Document(test_file)
    print(f"文档读取成功！")
    print(f"段落数量: {len(doc.paragraphs)}")
    
    styles = set()
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name:
            styles.add(para.style.name)
        if i < 5:  # 只打印前5个段落的样式
            print(f"  段落{i}: style={para.style.name if para.style else 'None'}")
    
    print(f"\n检测到的样式数量: {len(styles)}")
    print(f"样式列表: {sorted(list(styles))}")
    
except Exception as e:
    print(f"读取失败: {e}")
    import traceback
    traceback.print_exc()
