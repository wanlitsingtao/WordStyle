# -*- coding: utf-8 -*-
"""
测试 8js3.docx 的转换
"""
import os
from doc_converter import DocumentConverter
from docx import Document

def test_8js3_conversion():
    source_file = "e:\\LingMa\\8js3.docx"
    template_file = "e:\\LingMa\\mb.docx"
    output_file = "e:\\LingMa\\8js3_test_output.docx"
    
    print("=" * 60)
    print("开始测试 8js3.docx 转换")
    print("=" * 60)
    
    if not os.path.exists(source_file):
        print(f"错误：源文件不存在 {source_file}")
        return
    
    if not os.path.exists(template_file):
        print(f"错误：模板文件不存在 {template_file}")
        return
    
    # 执行转换
    print("\n正在转换...")
    converter = DocumentConverter()
    success, message = converter.convert_styles(source_file, template_file, output_file)
    
    print(f"\n转换结果: {'成功' if success else '失败'}")
    print(f"消息: {message}")
    
    if success and os.path.exists(output_file):
        print(f"\n✓ 输出文件已创建: {output_file}")
        
        # 检查输出文件中的特殊对象
        print("\n正在检查输出文件中的 Visio 图...")
        try:
            output_doc = Document(output_file)
            
            object_count = 0
            shape_count = 0
            para_with_objects = 0
            
            for idx, para in enumerate(output_doc.paragraphs):
                objects = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
                shapes = para._element.findall('.//{urn:schemas-microsoft-com:vml}shape')
                
                if objects or shapes:
                    para_with_objects += 1
                    object_count += len(objects)
                    shape_count += len(shapes)
            
            print(f"\n输出文件统计:")
            print(f"  总段落数: {len(output_doc.paragraphs)}")
            print(f"  包含特殊对象的段落数: {para_with_objects}")
            print(f"  OLE 对象总数: {object_count}")
            print(f"  VML 形状总数: {shape_count}")
            
            if object_count > 0 or shape_count > 0:
                print(f"\n✓ 成功复制了 Visio 图和 OLE 对象!")
                if para_with_objects >= 60:  # 期望至少复制大部分对象
                    print(f"✓ 复制率良好 ({para_with_objects}/69)")
                else:
                    print(f"⚠ 复制率较低 ({para_with_objects}/69)")
            else:
                print(f"\n✗ 未能复制任何 Visio 图或 OLE 对象")
        
        except Exception as e:
            print(f"✗ 检查输出文件时出错: {e}")
            import traceback
            traceback.print_exc()
    else:
        print(f"\n✗ 输出文件未创建")
    
    print("\n" + "=" * 60)
    print("测试完成")
    print("=" * 60)

if __name__ == "__main__":
    test_8js3_conversion()
