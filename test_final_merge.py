# -*- coding: utf-8 -*-
"""
测试完整合并单元格功能(横向+纵向)
"""
from doc_converter import DocumentConverter
import os
import time

print("=" * 60)
print("测试完整合并单元格功能")
print("=" * 60)

converter = DocumentConverter()

source_file = "8js3.docx"
template_file = "mb.docx"
output_file = "test_full_merge_result.docx"

print(f"\n源文件: {source_file}")
print(f"模板文件: {template_file}")
print(f"输出文件: {output_file}")
print(f"\n开始转换...")
print("-" * 60)

start_time = time.time()

success, msg = converter.convert_styles(source_file, template_file, output_file)

elapsed_time = time.time() - start_time

print("-" * 60)
if success:
    print(f"✅ 转换成功!")
    print(f"⏱️  耗时: {elapsed_time:.2f} 秒")
    
    if os.path.exists(output_file):
        file_size = os.path.getsize(output_file)
        print(f"📄 文件大小: {file_size / 1024 / 1024:.2f} MB")
        print(f"📍 文件路径: {os.path.abspath(output_file)}")
        
        # 验证表格数量
        from docx import Document
        src_doc = Document(source_file)
        tgt_doc = Document(output_file)
        print(f"\n📊 表格数量验证:")
        print(f"   源文档: {len(src_doc.tables)} 个表格")
        print(f"   目标文档: {len(tgt_doc.tables)} 个表格")
        if len(src_doc.tables) == len(tgt_doc.tables):
            print(f"   ✅ 表格数量一致")
        else:
            print(f"   ❌ 表格数量不一致!")
    else:
        print("❌ 文件未生成")
else:
    print(f"❌ 转换失败: {msg}")

print("\n" + "=" * 60)
