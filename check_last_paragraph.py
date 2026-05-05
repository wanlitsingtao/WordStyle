"""检查文档最后一段的情况"""
from docx import Document
from docx.oxml.ns import qn

doc = Document("test_after_heading_result_1919.docx")

print("=" * 80)
print("文档最后10段详情")
print("=" * 80)

for i in range(max(0, len(doc.paragraphs)-10), len(doc.paragraphs)):
    para = doc.paragraphs[i]
    style_name = para.style.name if para.style else "None"
    text = para.text[:60]
    marker = "[应答]" if "应答" in text else ""
    
    # 检查是否有大纲级别
    p_pr = para._element.find(qn('w:pPr'))
    has_outline = False
    if p_pr is not None:
        outline_elem = p_pr.find(qn('w:outlineLvl'))
        has_outline = outline_elem is not None
    
    print(f"{i}: [{style_name:20s}] {marker:8s} outline={has_outline} | {text}")

print("\n" + "=" * 80)
print("检查最后一个非空段落")
print("=" * 80)

# 从后往前找第一个非空段落
for i in range(len(doc.paragraphs)-1, -1, -1):
    para = doc.paragraphs[i]
    if para.text.strip():  # 非空
        style_name = para.style.name if para.style else "None"
        text = para.text[:60]
        
        # 检查是否为标题
        p_pr = para._element.find(qn('w:pPr'))
        is_heading = False
        if p_pr is not None:
            outline_elem = p_pr.find(qn('w:outlineLvl'))
            if outline_elem is not None:
                is_heading = True
        
        print(f"位置 {i}: [{style_name}] heading={is_heading}")
        print(f"内容: {text}")
        break
