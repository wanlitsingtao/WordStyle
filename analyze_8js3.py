# -*- coding: utf-8 -*-
"""
分析 8js3.docx 中的 Visio 图
"""
from docx import Document

def analyze_8js3():
    source_file = "e:\\LingMa\\8js3.docx"
    
    print("正在分析 8js3.docx...")
    print("=" * 60)
    
    try:
        doc = Document(source_file)
        
        object_count = 0
        shape_count = 0
        para_with_objects = []
        
        for idx, para in enumerate(doc.paragraphs):
            objects = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
            shapes = para._element.findall('.//{urn:schemas-microsoft-com:vml}shape')
            
            if objects or shapes:
                obj_count = len(objects)
                sha_count = len(shapes)
                object_count += obj_count
                shape_count += sha_count
                para_with_objects.append((idx, obj_count, sha_count))
                print(f"段落 {idx}: {obj_count} 个 OLE 对象, {sha_count} 个 VML 形状")
        
        print(f"\n统计结果:")
        print(f"  总段落数: {len(doc.paragraphs)}")
        print(f"  包含特殊对象的段落数: {len(para_with_objects)}")
        print(f"  OLE 对象总数: {object_count}")
        print(f"  VML 形状总数: {shape_count}")
        
        if para_with_objects:
            print(f"\n前10个包含对象的段落:")
            for i, (idx, obj, sha) in enumerate(para_with_objects[:10]):
                print(f"  {i+1}. 段落 {idx}: {obj} OLE, {sha} VML")
        
        # 检查表格中的对象
        table_object_count = 0
        table_shape_count = 0
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        objects = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
                        shapes = para._element.findall('.//{urn:schemas-microsoft-com:vml}shape')
                        if objects or shapes:
                            table_object_count += len(objects)
                            table_shape_count += len(shapes)
        
        print(f"\n表格中的对象:")
        print(f"  OLE 对象: {table_object_count}")
        print(f"  VML 形状: {table_shape_count}")
        
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    analyze_8js3()
