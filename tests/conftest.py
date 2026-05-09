# -*- coding: utf-8 -*-
"""
pytest 配置文件
提供全局fixture和测试配置
"""
import pytest
import os
import sys
from pathlib import Path

# 添加项目根目录到路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))


@pytest.fixture(scope="session")
def base_url():
    """后端API基础URL"""
    return os.getenv("BACKEND_URL", "http://localhost:8000")


@pytest.fixture(scope="session")
def frontend_url():
    """前端Streamlit应用URL"""
    return os.getenv("FRONTEND_URL", "http://localhost:8501")


@pytest.fixture(scope="session")
def test_files_dir():
    """测试文件目录"""
    test_dir = project_root / "tests" / "test_files"
    test_dir.mkdir(parents=True, exist_ok=True)
    return str(test_dir)


@pytest.fixture(scope="session")
def sample_docx_file(test_files_dir):
    """创建示例DOCX测试文件"""
    from docx import Document
    
    file_path = os.path.join(test_files_dir, "sample_test.docx")
    
    # 如果文件已存在，直接返回
    if os.path.exists(file_path):
        return file_path
    
    # 创建测试文档
    doc = Document()
    doc.add_heading('第一章 测试文档', level=1)
    doc.add_paragraph('这是一个用于自动化测试的示例文档。')
    doc.add_heading('1.1 测试章节', level=2)
    doc.add_paragraph('本段落用于验证文档转换功能。')
    
    # 添加表格
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '序号'
    table.cell(0, 1).text = '内容'
    table.cell(1, 0).text = '1'
    table.cell(1, 1).text = '测试项1'
    table.cell(2, 0).text = '2'
    table.cell(2, 1).text = '测试项2'
    
    doc.save(file_path)
    return file_path


@pytest.fixture(scope="session")
def template_docx_file(test_files_dir):
    """创建模板DOCX测试文件"""
    from docx import Document
    
    file_path = os.path.join(test_files_dir, "template_test.docx")
    
    # 如果文件已存在，直接返回
    if os.path.exists(file_path):
        return file_path
    
    # 创建模板文档
    doc = Document()
    doc.add_heading('模板标题', level=1)
    doc.add_heading('模板二级标题', level=2)
    para = doc.add_paragraph('模板正文样式')
    para.style = 'Normal'
    
    doc.save(file_path)
    return file_path
