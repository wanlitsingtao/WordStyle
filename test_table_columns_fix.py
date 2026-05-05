# -*- coding: utf-8 -*-
"""测试表格列数计算修复"""
from doc_converter import DocumentConverter

print("="*80)
print("测试表格列数计算修复")
print("="*80)

# 创建转换器
converter = DocumentConverter()

# 测试文件
source_file = "E:/LingMa/8js3.docx"
template_file = "E:/LingMa/mb.docx"
output_file = "E:/LingMa/8js3_test_merge.docx"

print(f"\n源文件: {source_file}")
print(f"模板文件: {template_file}")
print(f"输出文件: {output_file}")
print()

# 只转换表格 #36 来测试
# 但 full_convert 会转换整个文档，所以我们需要直接测试 copy_table_with_images

from docx import Document
from docx.oxml.ns import qn

# 打开源文档
source_doc = Document(source_file)
table_36 = source_doc.tables[35]  # 索引从0开始

print(f"测试表格 #36:")
print(f"源表格行数: {len(table_36.rows)}")

# 计算正确的列数
max_cols = 0
for row in table_36.rows:
    col_count = 0
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        span = 1
        if tcPr is not None:
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is not None:
                span = int(gridSpan.get(qn('w:val'), 1))
        col_count += span
    if col_count > max_cols:
        max_cols = col_count

print(f"计算得到的列数: {max_cols}")
print(f"python-docx 的 len(table.columns): {len(table_36.columns)}")
print()

# 现在执行完整转换
print("执行完整转换...")
try:
    success, actual_file, msg = converter.full_convert(
        source_file=source_file,
        template_file=template_file,
        output_file=output_file,
        do_mood=False,
        do_answer_insertion=False,
        progress_callback=lambda step, msg: print(f"  步骤 {step}: {msg}")
    )
    
    if success:
        print(f"\n✅ 转换成功！")
        print(f"实际输出文件: {actual_file}")
        
        # 分析转换后的表格 #36
        print("\n" + "="*80)
        print("分析转换后的表格 #36")
        print("="*80)
        
        converted_doc = Document(actual_file)
        converted_table_36 = converted_doc.tables[35]
        
        print(f"\n转换后的表格 #36:")
        print(f"行数: {len(converted_table_36.rows)}")
        print(f"列数: {len(converted_table_36.columns)}")
        print()
        
        # 检查前3行
        for i, row in enumerate(converted_table_36.rows[:3]):
            print(f"行 {i+1}: {len(row.cells)} 个 cells")
            for j, cell in enumerate(row.cells):
                text = ''.join([p.text for p in cell.paragraphs]).strip()
                if text:
                    print(f"  [{i},{j}]: '{text[:20]}'")
        
        print("\n" + "="*80)
        print("请打开文档检查表格是否正确！")
        print("="*80)
    else:
        print(f"\n❌ 转换失败: {msg}")

except Exception as e:
    print(f"\n❌ 发生错误: {e}")
    import traceback
    traceback.print_exc()
