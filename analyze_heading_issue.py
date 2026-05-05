"""分析2级标题和3级标题之间的问题"""
from docx import Document

def is_heading_by_style(para):
    """通过样式判断是否为标题"""
    if not para.style:
        return False
    style_name = para.style.name
    # 检查样式名称是否包含'heading'
    if 'heading' in style_name.lower():
        return True
    # 检查样式名称是否为纯数字
    if style_name.isdigit():
        return True
    return False

doc = Document("test_after_heading_result.docx")

print("=" * 80)
print("查找两个标题之间的应答句")
print("=" * 80)

error_count = 0
for i, para in enumerate(doc.paragraphs):
    if "应答" in para.text:
        # 检查前后段落
        prev_is_heading = False
        next_is_heading = False
        prev_text = ""
        next_text = ""
        
        if i > 0:
            prev_para = doc.paragraphs[i-1]
            prev_style = prev_para.style.name if prev_para.style else "None"
            prev_text = prev_para.text[:40]
            prev_is_heading = is_heading_by_style(prev_para)
        
        if i + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i+1]
            next_style = next_para.style.name if next_para.style else "None"
            next_text = next_para.text[:40]
            next_is_heading = is_heading_by_style(next_para)
        
        # 如果前后都是标题，说明有问题
        if prev_is_heading and next_is_heading:
            error_count += 1
            print(f"\n错误位置 {i}:")
            print(f"   前一段 [标题({prev_style})]: {prev_text}")
            print(f"   [应答句]: {para.text[:40]}")
            print(f"   后一段 [标题({next_style})]: {next_text}")

print(f"\n\n总共发现 {error_count} 个错误")

# 检查文档最后一段
print("\n" + "=" * 80)
print("检查文档最后5段")
print("=" * 80)

for i in range(max(0, len(doc.paragraphs)-5), len(doc.paragraphs)):
    para = doc.paragraphs[i]
    style_name = para.style.name if para.style else "None"
    text = para.text[:60]
    marker = "[应答]" if "应答" in text else ""
    print(f"{i}: [{style_name:20s}] {marker:8s} {text}")
